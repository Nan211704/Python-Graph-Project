import os
import json
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend suitable for Flask servers
import matplotlib.pyplot as plt
from flask import Blueprint, request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from bson.objectid import ObjectId
from datetime import datetime 
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
import openpyxl
import tempfile
import re
import zipfile
import shutil
import plotly.graph_objects as go
from plotly.subplots import make_subplots


import re

# Define a constant for the section1_chart attribut

# Define a custom UPLOAD_FOLDER for this blueprint or ensure it's globally configured
UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'uploads')
# Ensure the upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'csv', 'xlsx', 'docx'}
ALLOWED_REPORT_EXTENSIONS = {'csv', 'xlsx'}

projects_bp = Blueprint('projects', __name__)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_report_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_REPORT_EXTENSIONS

def extract_report_info_from_excel(excel_path):
    """Extract Report_Name and Report_Code from Excel file"""
    try:
        # Load the Excel file
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        sheet = wb.active
        
        report_name = None
        report_code = None
        
        # Search for the columns in the first few rows
        for row_idx in range(1, min(10, sheet.max_row + 1)):  # Check first 10 rows
            for col_idx in range(1, min(10, sheet.max_column + 1)):  # Check first 10 columns
                cell_value = sheet.cell(row=row_idx, column=col_idx).value
                
                if cell_value and isinstance(cell_value, str):
                    cell_value = cell_value.strip()
                    
                    # Look for Report_Name column
                    if cell_value.lower() == 'report_name':
                        # Get the value from the next row in the same column
                        if row_idx + 1 <= sheet.max_row:
                            report_name = sheet.cell(row=row_idx + 1, column=col_idx).value
                            if report_name:
                                report_name = str(report_name).strip()
                    
                    # Look for Report_Code column
                    elif cell_value.lower() == 'report_code':
                        # Get the value from the next row in the same column
                        if row_idx + 1 <= sheet.max_row:
                            report_code = sheet.cell(row=row_idx + 1, column=col_idx).value
                            if report_code:
                                report_code = str(report_code).strip()
        
        wb.close()
        
        # Fallback to filename if not found
        if not report_name:
            report_name = os.path.splitext(os.path.basename(excel_path))[0]
            current_app.logger.warning(f"Report_Name not found in {excel_path}, using filename: {report_name}")
        
        if not report_code:
            report_code = f"REPORT_{os.path.splitext(os.path.basename(excel_path))[0]}"
            current_app.logger.warning(f"Report_Code not found in {excel_path}, using generated code: {report_code}")
        
        current_app.logger.info(f"Extracted from {excel_path}: Report_Name='{report_name}', Report_Code='{report_code}'")
        return report_name, report_code
        
    except Exception as e:
        current_app.logger.error(f"Error extracting report info from {excel_path}: {e}")
        # Fallback to filename-based naming
        fallback_name = os.path.splitext(os.path.basename(excel_path))[0]
        fallback_code = f"REPORT_{fallback_name}"
        return fallback_name, fallback_code

def create_expanded_pie_chart(labels, values, colors, expanded_segment, title, value_format=""):
    """
    Create an expanded pie chart with one segment shown as a bar chart
    """
    fig = make_subplots(
        rows=1, cols=2,
        specs=[[{"type": "pie"}, {"type": "bar"}]],
        subplot_titles=(title, f"{expanded_segment} Details")
    )
    
    # Add pie chart
    fig.add_trace(go.Pie(
        labels=labels,
        values=values,
        textinfo="label+percent",
        textposition="outside",
        marker=dict(colors=colors)
    ), row=1, col=1)
    
    # Add bar chart for expanded segment
    if expanded_segment in labels:
        segment_idx = labels.index(expanded_segment)
        segment_value = values[segment_idx]
        segment_color = colors[segment_idx] if segment_idx < len(colors) else colors[0]
        
        fig.add_trace(go.Bar(
            x=[expanded_segment],
            y=[segment_value],
            marker_color=segment_color,
            text=[f"{segment_value}{value_format}"],
            textposition="auto"
        ), row=1, col=2)
    
    fig.update_layout(
        title_text=title,
        showlegend=False,
        height=500
    )
    
    return fig

def create_bar_of_pie_chart(labels, values, other_labels, other_values, colors, other_colors, title, value_format=""):
    """
    Create a 'bar of pie' chart using Plotly: pie chart with one segment broken down as a bar chart.
    """
    # Filter out empty/null values from other_labels and other_values
    filtered_data = []
    for label, value in zip(other_labels, other_values):
        # Check if both label and value are not empty/null
        if (label is not None and str(label).strip() != "" and 
            value is not None and str(value).strip() != "" and 
            str(value).strip() != "0"):
            filtered_data.append((label, value))
    
    if filtered_data:
        filtered_labels, filtered_values = zip(*filtered_data)
    else:
        # If no valid data, use empty lists
        filtered_labels, filtered_values = [], []
    
    fig = make_subplots(
        rows=1, cols=2,
        specs=[[{"type": "pie"}, {"type": "bar"}]],
        column_widths=[0.5, 0.5],
        subplot_titles=(title, "Breakdown of 'Other'")
    )
    # Main pie chart
    fig.add_trace(go.Pie(
        labels=labels,
        values=values,
        marker=dict(colors=colors),
        textinfo="percent",
        hoverinfo="label+percent+value",
        pull=[0.1 if l == "Other" else 0 for l in labels],
        name="Main Pie"
    ), row=1, col=1)
    
    # Bar chart for breakdown (only if we have filtered data)
    if filtered_labels and filtered_values:
        # Create individual bar traces for each data point to avoid stacking
        for i, (label, value) in enumerate(zip(filtered_labels, filtered_values)):
            # Use individual colors for each bar
            bar_color = other_colors[i] if other_colors and i < len(other_colors) else None
            
    fig.add_trace(go.Bar(
                x=[label],  # Single x value for each bar
                y=[value],  # Single y value for each bar
                marker_color=bar_color,
                text=[f"{value}{value_format}"],
        textposition="auto",
                name=f"Breakdown {i+1}",
                showlegend=False  # Hide individual bar legends
    ), row=1, col=2)
    
    fig.update_layout(
        title_text=title,
        showlegend=False,
        height=500,
        width=900
    )
    return fig

def _generate_report(project_id, template_path, data_file_path):
    import pandas as pd
    import json
    import tempfile
    import matplotlib.pyplot as plt
    from matplotlib.ticker import FuncFormatter
    from docx import Document
    from docx.shared import Inches
    import re
    import os

    plt.style.use('ggplot')  # 👈 Apply a cleaner visual style

    try:
        # Report generation started

        # Try to read Excel with original formatting preserved
        try:
            df = pd.read_excel(data_file_path, sheet_name=0, keep_default_na=False)  # Use first sheet
        except:
            df = pd.read_excel(data_file_path, sheet_name=0)  # Fallback to default
        
        df.columns = df.columns.str.strip().str.replace(" ", "_").str.replace("__", "_")
        
        # Log the raw data to see what pandas is reading
        # Excel data loaded successfully

        # Excel structure loaded silently
        
        # Validate required columns exist
        required_columns = ["Text_Tag", "Text", "Chart_Tag", "Chart_Attributes", "Chart_Type"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            current_app.logger.error(f"❌ Missing required columns: {missing_columns}")
            raise ValueError(f"Excel file missing required columns: {missing_columns}")

        text_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Text_Tag"], df["Text"]) if pd.notna(k) and pd.notna(v)}
        chart_attr_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Chart_Tag"], df["Chart_Attributes"]) if pd.notna(k) and pd.notna(v)}
        chart_type_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Chart_Tag"], df["Chart_Type"]) if pd.notna(k) and pd.notna(v)}
        
        # Data maps created silently

        flat_data_map = {}
        # Add global metadata columns that can be used across all sections
        global_metadata = {}
        
        # First pass: Process global metadata columns from ALL rows
        for _, row in df.iterrows():
            for col in df.columns:
                col_lower = col.lower().strip()
                
                # Handle global metadata columns (can be used across all sections)
                if col_lower in ['country', 'report_name', 'report_code', 'currency']:
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Store in global metadata for use across all sections
                        global_metadata[col_lower] = str(value).strip()
                        # Also add to flat_data_map with the column name as key
                        flat_data_map[col_lower] = str(value).strip()
                        # Global metadata processed
                    else:
                        current_app.logger.warning(f"⚠️ Empty or null value for global metadata column: {col} (value: {value})")
        
        # Ensure we have all required global metadata - if any are missing, log a warning
        required_global_metadata = ['country', 'report_name', 'report_code', 'currency']
        missing_metadata = [key for key in required_global_metadata if key not in flat_data_map]
        if missing_metadata:
            current_app.logger.warning(f"⚠️ Missing global metadata: {missing_metadata}")
            current_app.logger.info(f"📋 Available global metadata: {list(flat_data_map.keys())}")
        else:
            current_app.logger.info(f"✅ All required global metadata found: {list(flat_data_map.keys())}")
        
        # Second pass: Process chart-specific data from rows with chart tags
        for _, row in df.iterrows():
            chart_tag = row.get("Chart_Tag")
            if not isinstance(chart_tag, str) or not chart_tag:
                continue
            section_prefix = chart_tag.replace('_chart', '').lower()

            # Process all columns systematically
            for col in df.columns:
                col_lower = col.lower().strip()
                
                # Skip global metadata columns as they're already processed
                if col_lower in ['country', 'report_name', 'report_code', 'currency']:
                    continue
                
                # Handle Chart Data columns (for values)
                if col_lower.startswith("chart_data_y"):
                    year = col.replace("Chart_Data_", "").replace("chart_data_", "")
                    key = f"{section_prefix}_{year.lower()}"
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Format values to always show one decimal place
                        try:
                            float_val = float(value)
                            formatted_val = f"{float_val:.1f}"
                            flat_data_map[key] = formatted_val
                        except (ValueError, TypeError):
                            # If conversion fails, use the original value as string
                            flat_data_map[key] = str(value).strip()
                    else:
                        current_app.logger.warning(f"⚠️ Empty value for {col} -> {key}")
                
                # Handle Growth columns (for growth rates)
                elif col_lower.startswith("growth_y"):
                    year = col.replace("Growth_", "").replace("growth_", "")
                    key = f"{section_prefix}_{year.lower()}_kpi2"
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Convert percentage values from Excel to decimal format for table display
                        try:
                            # Log the raw value and its type
                            # Processing growth value
                            
                            # Convert to float first to handle any numeric format
                            float_val = float(value)
                            # Growth value processed
                            
                            # Check if the original value string contains '%' (Excel percentage format)
                            original_str = str(value).strip()
                            if '%' in original_str:
                                # Remove % and convert to decimal
                                clean_val = original_str.replace('%', '').strip()
                                float_val = float(clean_val)
                                decimal_val = f"{float_val / 100:.3f}"
                                flat_data_map[key] = decimal_val
                                # Excel percentage converted
                            elif float_val > 1:
                                # Convert percentage to decimal (e.g., 20% -> 0.2)
                                decimal_val = f"{float_val / 100:.3f}"
                                flat_data_map[key] = decimal_val
                                # Percentage converted to decimal
                            else:
                                # If it's already a decimal, format it properly
                                decimal_val = f"{float_val:.3f}"
                                flat_data_map[key] = decimal_val
                                # Decimal formatted
                        except (ValueError, TypeError):
                            # If conversion fails, use the original value as string
                            flat_data_map[key] = str(value).strip()
                            current_app.logger.warning(f"Could not convert growth value: {value}")
                    else:
                        current_app.logger.warning(f"⚠️ Empty value for {col} -> {key}")

            # Handle CAGR
            if pd.notna(row.get("Chart_Data_CAGR")):
                key = f"{section_prefix}_cgrp"
                value = row["Chart_Data_CAGR"]
                if str(value).strip():
                    # Convert CAGR percentage values from Excel to decimal format for table display
                    try:
                        # Convert to float first to handle any numeric format
                        float_val = float(value)
                        # CAGR value processed
                        
                        # Check if the original value string contains '%' (Excel percentage format)
                        original_str = str(value).strip()
                        if '%' in original_str:
                            # Remove % and convert to decimal
                            clean_val = original_str.replace('%', '').strip()
                            float_val = float(clean_val)
                            decimal_val = f"{float_val / 100:.3f}"
                            flat_data_map[key] = decimal_val
                            # Excel CAGR percentage converted
                        elif float_val > 1:
                            # Convert percentage to decimal (e.g., 10.5% -> 0.105)
                            decimal_val = f"{float_val / 100:.3f}"
                            flat_data_map[key] = decimal_val
                            # CAGR percentage converted to decimal
                        else:
                            # If it's already a decimal, format it properly
                            decimal_val = f"{float_val:.3f}"
                            flat_data_map[key] = decimal_val
                            # CAGR decimal formatted
                    except (ValueError, TypeError):
                        # If conversion fails, use the original value as string
                        flat_data_map[key] = str(value).strip()
                        current_app.logger.warning(f"Could not convert CAGR value: {value}")
                else:
                    current_app.logger.warning(f"⚠️ Empty value for Chart_Data_CAGR -> {key}")

        # Ensure all keys are lowercase
        flat_data_map = {k.lower(): v for k, v in flat_data_map.items()}
        
        # Log the final data map for debugging
        # Data map prepared for replacement
        
        # Data mapping completed silently

        doc = Document(template_path)

        def replace_text_in_paragraph(paragraph):
            nonlocal flat_data_map, text_map  # Access variables from outer scope
            # First, try to replace placeholders that are contained within single runs
            for run in paragraph.runs:
                original_text = run.text
                
                # Handle ${} format placeholders
                matches = re.findall(r"\$\{(.*?)\}", run.text)
                for match in matches:
                    key_lower = match.lower().strip()
                    val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    if val:
                        pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                        old_text = run.text
                        run.text = pattern.sub(val, run.text)
                        # Log replacements for growth rates, CAGR, and global metadata
                        if 'kpi2' in key_lower or 'cgrp' in key_lower or key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.info(f"🔄 TABLE REPLACEMENT: ${{{match}}} -> {val} (in text: {old_text} -> {run.text})")
                    else:
                        current_app.logger.warning(f"⚠️ No value found for placeholder ${{{match}}} (key: {key_lower})")
                        # Enhanced debugging for global metadata
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ CRITICAL: Global metadata '{key_lower}' not found in flat_data_map!")
                            current_app.logger.error(f"❌ Available keys in flat_data_map: {list(flat_data_map.keys())}")
                            current_app.logger.error(f"❌ Available keys in text_map: {list(text_map.keys())}")
                
                # Handle <> format placeholders
                angle_matches = re.findall(r"<(.*?)>", run.text)
                for match in angle_matches:
                    key_lower = match.lower().strip()
                    val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    if val:
                        pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                        old_text = run.text
                        run.text = pattern.sub(val, run.text)
                        # Log replacements for global metadata
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.info(f"🔄 ANGLE REPLACEMENT: <{match}> -> {val} (in text: {old_text} -> {run.text})")
                    else:
                        current_app.logger.warning(f"⚠️ No value found for placeholder <{match}> (key: {key_lower})")
                        # Enhanced debugging for global metadata
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ CRITICAL: Global metadata '{key_lower}' not found in flat_data_map!")
                            current_app.logger.error(f"❌ Available keys in flat_data_map: {list(flat_data_map.keys())}")
                            current_app.logger.error(f"❌ Available keys in text_map: {list(text_map.keys())}")
            
            # Then, handle placeholders that might be split across multiple runs
            # Get the full paragraph text to find all placeholders
            full_text = paragraph.text
            
            # Handle ${} format placeholders across multiple runs
            all_matches = re.findall(r"\$\{(.*?)\}", full_text)
            for match in all_matches:
                key_lower = match.lower().strip()
                val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                if val:
                    placeholder = f"${{{match}}}"
                    if placeholder in full_text:
                        # Log replacements for growth rates, CAGR, and global metadata
                        if 'kpi2' in key_lower or 'cgrp' in key_lower or key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.info(f"🔄 TABLE REPLACEMENT (multi-run): ${{{match}}} -> {val}")
                        
                        # Clear all runs and recreate with replacement
                        for run in paragraph.runs:
                            run.text = ""
                        # Set the first run to the replaced text
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text.replace(placeholder, str(val))
                        else:
                            # If no runs exist, create one
                            paragraph.add_run(full_text.replace(placeholder, str(val)))
                        break  # Only process one placeholder at a time to avoid conflicts
            
            # Handle <> format placeholders across multiple runs
            all_angle_matches = re.findall(r"<(.*?)>", full_text)
            for match in all_angle_matches:
                key_lower = match.lower().strip()
                val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                if val:
                    placeholder = f"<{match}>"
                    if placeholder in full_text:
                        # Log replacements for global metadata
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.info(f"🔄 ANGLE REPLACEMENT (multi-run): <{match}> -> {val}")
                        
                        # Clear all runs and recreate with replacement
                        for run in paragraph.runs:
                            run.text = ""
                        # Set the first run to the replaced text
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text.replace(placeholder, str(val))
                        else:
                            # If no runs exist, create one
                            paragraph.add_run(full_text.replace(placeholder, str(val)))
                        break  # Only process one placeholder at a time to avoid conflicts

        def replace_text_in_tables():
            nonlocal doc, flat_data_map, text_map  # Access variables from outer scope
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_text_in_paragraph(para)
            
            # Table processing completed silently

        def process_entire_document():
            """Process the entire document comprehensively to catch all placeholders"""
            nonlocal doc, flat_data_map, text_map  # Access variables from outer scope
            # current_app.logger.info("🔄 PROCESSING ENTIRE DOCUMENT COMPREHENSIVELY")
            
            # Find ALL placeholders in the entire document first
            # current_app.logger.info("🔍 SEARCHING FOR ALL PLACEHOLDERS IN ENTIRE DOCUMENT")
            
            all_placeholders_found = set()
            
            # Search through ALL paragraphs, tables, headers, footers, etc.
            def search_for_placeholders(container):
                """Search for placeholders in any container (document, table, header, footer)"""
                if hasattr(container, 'paragraphs'):
                    for para in container.paragraphs:
                        if para.text:
                            current_app.logger.debug(f"🔍 Searching paragraph: '{para.text}'")
                            # Find ${} placeholders
                            dollar_matches = re.findall(r"\$\{(.*?)\}", para.text)
                            for match in dollar_matches:
                                all_placeholders_found.add(f"${{{match}}}")
                                current_app.logger.debug(f"  ✅ Found $ placeholder: ${{{match}}}")
                            
                            # Find <> placeholders
                            angle_matches = re.findall(r"<(.*?)>", para.text)
                            for match in angle_matches:
                                all_placeholders_found.add(f"<{match}>")
                                current_app.logger.debug(f"  ✅ Found <> placeholder: <{match}>")
                
                if hasattr(container, 'tables'):
                    for table in container.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                search_for_placeholders(cell)
            
            # Search in main document
            search_for_placeholders(doc)
            
            # Search in headers and footers
            for section in doc.sections:
                if section.header:
                    search_for_placeholders(section.header)
                if section.footer:
                    search_for_placeholders(section.footer)
            
            # Additional search: Look at raw XML for any missed placeholders
            # current_app.logger.info("🔍 ADDITIONAL SEARCH: Looking at raw XML for missed placeholders")
            try:
                for element in doc.element.iter():
                    if hasattr(element, 'text') and element.text:
                        # Find ${} placeholders
                        dollar_matches = re.findall(r"\$\{(.*?)\}", element.text)
                        for match in dollar_matches:
                            all_placeholders_found.add(f"${{{match}}}")
                            current_app.logger.debug(f"  ✅ Found $ placeholder in XML: ${{{match}}}")
                        
                        # Find <> placeholders
                        angle_matches = re.findall(r"<(.*?)>", element.text)
                        for match in angle_matches:
                            all_placeholders_found.add(f"<{match}>")
                            current_app.logger.debug(f"  ✅ Found <> placeholder in XML: <{match}>")
            except Exception as e:
                current_app.logger.warning(f"⚠️ Error in additional XML search: {e}")
            
            # current_app.logger.info(f"🔍 Found {len(all_placeholders_found)} unique placeholders: {list(all_placeholders_found)}")
            
            # Log specific global metadata placeholders found
            global_placeholders_found = [p for p in all_placeholders_found if any(key in p.lower() for key in ['country', 'report_name', 'report_code', 'currency'])]
            # current_app.logger.info(f"🔍 Global metadata placeholders found: {global_placeholders_found}")
            
            # Verify that we have data for all found global metadata placeholders
            for placeholder in global_placeholders_found:
                if placeholder.startswith('${'):
                    key = placeholder[2:-1].lower()
                elif placeholder.startswith('<') and placeholder.endswith('>'):
                    key = placeholder[1:-1].lower()
                else:
                    continue
                
                if key in ['country', 'report_name', 'report_code', 'currency']:
                      if key in flat_data_map:
                          current_app.logger.info(f"✅ Data available for {placeholder}: {flat_data_map[key]}")
                      else:
                          current_app.logger.error(f"❌ NO DATA AVAILABLE for {placeholder} (key: {key})")
                          current_app.logger.error(f"❌ Available keys: {list(flat_data_map.keys())}")
            
            # Now replace ALL placeholders everywhere they appear
            current_app.logger.info("🔄 REPLACING ALL PLACEHOLDERS EVERYWHERE")
            
            # Multiple passes to ensure we catch everything
            for pass_num in range(3):
                current_app.logger.info(f"🔄 Pass {pass_num + 1}/3: Processing all document elements...")
                
                # Process main document
                for para in doc.paragraphs:
                    replace_text_in_paragraph(para)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_text_in_paragraph(para)
                
                # Process headers and footers
                for section in doc.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            replace_text_in_paragraph(para)
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        replace_text_in_paragraph(para)
                    
                    if section.footer:
                        for para in section.footer.paragraphs:
                            replace_text_in_paragraph(para)
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        replace_text_in_paragraph(para)
                
                # Process XML elements that might contain placeholders (like cover page content)
                current_app.logger.info(f"🔄 Pass {pass_num + 1}/3: Processing XML elements for special content...")
                try:
                    for element in doc.element.iter():
                        if hasattr(element, 'text') and element.text:
                            # Check if this element contains any placeholders we need to replace
                            original_text = element.text
                            modified_text = original_text
                            
                            # Replace ${} placeholders
                            for placeholder in all_placeholders_found:
                                if placeholder.startswith('${'):
                                    key = placeholder[2:-1].lower()  # Remove ${} and convert to lowercase
                                    if key in flat_data_map:
                                        value = flat_data_map[key]
                                        modified_text = modified_text.replace(placeholder, str(value))
                                        current_app.logger.info(f"🔄 XML REPLACEMENT: {placeholder} -> {value} (in XML element)")
                            
                            # Replace <> placeholders
                            for placeholder in all_placeholders_found:
                                if placeholder.startswith('<') and placeholder.endswith('>'):
                                    key = placeholder[1:-1].lower()  # Remove <> and convert to lowercase
                                    if key in flat_data_map:
                                        value = flat_data_map[key]
                                        modified_text = modified_text.replace(placeholder, str(value))
                                        current_app.logger.info(f"🔄 XML REPLACEMENT: {placeholder} -> {value} (in XML element)")
                            
                            # Update the element text if it was modified
                            if modified_text != original_text:
                                element.text = modified_text
                                current_app.logger.info(f"🔄 XML ELEMENT UPDATED: '{original_text}' -> '{modified_text}'")
                except Exception as e:
                    #current_app.logger.warning(f"⚠️ Error processing XML elements: {e}")
            
            #current_app.logger.info("✅ COMPREHENSIVE DOCUMENT PROCESSING COMPLETED")
            
            # Additional pass: Handle special Word elements that might contain placeholders
                    current_app.logger.info("🔄 FINAL PASS: Processing special Word elements...")
            
            # Process text boxes and other special elements
            try:
                for shape in doc.inline_shapes:
                    if hasattr(shape, 'text_frame'):
                        for para in shape.text_frame.paragraphs:
                            replace_text_in_paragraph(para)
            except Exception as e:
                current_app.logger.warning(f"⚠️ Error processing inline shapes: {e}")
            
            # Process any remaining XML elements that might contain placeholders
            try:
                for element in doc.element.iter():
                    if hasattr(element, 'text') and element.text:
                        original_text = element.text
                        modified_text = original_text
                        
                        # Replace all known placeholders in this element
                        for key, value in flat_data_map.items():
                            if key in ['country', 'report_name', 'report_code', 'currency']:
                                # Try both ${} and <> formats
                                dollar_placeholder = f"${{{key}}}"
                                angle_placeholder = f"<{key}>"
                                
                                if dollar_placeholder in modified_text:
                                    modified_text = modified_text.replace(dollar_placeholder, str(value))
                                    #current_app.logger.info(f"🔄 FINAL XML REPLACEMENT: {dollar_placeholder} -> {value}")
                                
                                if angle_placeholder in modified_text:
                                    modified_text = modified_text.replace(angle_placeholder, str(value))
                                    #current_app.logger.info(f"🔄 FINAL XML REPLACEMENT: {angle_placeholder} -> {value}")
                        
                        # Update the element if modified
                        if modified_text != original_text:
                            element.text = modified_text
                            #current_app.logger.info(f"🔄 FINAL XML ELEMENT UPDATED: '{original_text}' -> '{modified_text}'")
            except Exception as e:
                current_app.logger.warning(f"⚠️ Error in final XML processing: {e}")
            
            current_app.logger.info("✅ FINAL PASS COMPLETED")
            
            # ULTIMATE PASS: Force replacement of all global metadata placeholders anywhere in the document
            #current_app.logger.info("🔄 ULTIMATE PASS: Force replacement of all global metadata placeholders...")
            
            # Get all global metadata values
            global_metadata_values = {
                'country': flat_data_map.get('country', ''),
                'report_name': flat_data_map.get('report_name', ''),
                'report_code': flat_data_map.get('report_code', ''),
                'currency': flat_data_map.get('currency', '')
            }
            
            #current_app.logger.info(f"📋 Global metadata values for ultimate replacement: {global_metadata_values}")
            
            # Force replacement in ALL possible locations
            def force_replace_in_element(element):
                """Force replace placeholders in any element that has text"""
                if hasattr(element, 'text') and element.text:
                    original_text = element.text
                    modified_text = original_text
                    
                    # Replace all possible formats of global metadata placeholders
                    for key, value in global_metadata_values.items():
                        if value:  # Only replace if we have a value
                            # Try all possible case variations
                            variations = [
                                f"${{{key}}}",
                                f"${{{key.upper()}}}",
                                f"${{{key.title()}}}",
                                f"<{key}>",
                                f"<{key.upper()}>",
                                f"<{key.title()}>",
                                f"<{key.replace('_', '')}>",
                                f"<{key.replace('_', '').upper()}>",
                                f"<{key.replace('_', '').title()}>"
                            ]
                            
                            for variation in variations:
                                if variation in modified_text:
                                    modified_text = modified_text.replace(variation, str(value))
                                    #current_app.logger.info(f"🔄 ULTIMATE REPLACEMENT: {variation} -> {value}")
                    
                    # Update the element if modified
                    if modified_text != original_text:
                        element.text = modified_text
                        #current_app.logger.info(f"🔄 ULTIMATE ELEMENT UPDATED: '{original_text}' -> '{modified_text}'")
            
            # Process ALL possible document elements
            try:
                # Process main document paragraphs
                for para in doc.paragraphs:
                    force_replace_in_element(para)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            force_replace_in_element(cell)
                            for para in cell.paragraphs:
                                force_replace_in_element(para)
                
                # Process headers and footers
                for section in doc.sections:
                    if section.header:
                        force_replace_in_element(section.header)
                        for para in section.header.paragraphs:
                            force_replace_in_element(para)
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    force_replace_in_element(cell)
                                    for para in cell.paragraphs:
                                        force_replace_in_element(para)
                    
                    if section.footer:
                        force_replace_in_element(section.footer)
                        for para in section.footer.paragraphs:
                            force_replace_in_element(para)
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    force_replace_in_element(cell)
                                    for para in cell.paragraphs:
                                        force_replace_in_element(para)
                
                # Process inline shapes (text boxes, etc.)
                for shape in doc.inline_shapes:
                    if hasattr(shape, 'text_frame'):
                        force_replace_in_element(shape.text_frame)
                        for para in shape.text_frame.paragraphs:
                            force_replace_in_element(para)
                
                # Process ALL XML elements recursively
                def process_xml_element(element):
                    """Recursively process all XML elements"""
                    force_replace_in_element(element)
                    
                    # Process child elements
                    for child in element:
                        process_xml_element(child)
                
                # Process the entire document XML
                process_xml_element(doc.element)
                
            except Exception as e:
                current_app.logger.error(f"❌ Error in ultimate pass: {e}")
                import traceback
                current_app.logger.error(f"❌ Traceback: {traceback.format_exc()}")
            
            #current_app.logger.info("✅ ULTIMATE PASS COMPLETED")
            
            # NUCLEAR OPTION: Direct XML manipulation to catch everything
            #current_app.logger.info("🔄 NUCLEAR OPTION: Direct XML manipulation...")
            
            try:
                # Convert document to XML string
                xml_content = doc._element.xml
                original_xml = xml_content
                
                # Replace all possible variations in the XML directly
                for key, value in global_metadata_values.items():
                    if value:
                        # Create all possible variations
                        variations = [
                            f"${{{key}}}",
                            f"${{{key.upper()}}}",
                            f"${{{key.title()}}}",
                            f"<{key}>",
                            f"<{key.upper()}>",
                            f"<{key.title()}>",
                            f"<{key.replace('_', '')}>",
                            f"<{key.replace('_', '').upper()}>",
                            f"<{key.replace('_', '').title()}>",
                            # Add more variations for common naming patterns
                            f"<{key.replace('_', ' ')}>",
                            f"<{key.replace('_', ' ').title()}>",
                            f"<{key.replace('_', ' ').upper()}>"
                        ]
                        
                        for variation in variations:
                            if variation in xml_content:
                                xml_content = xml_content.replace(variation, str(value))
                                #current_app.logger.info(f"🔄 NUCLEAR XML REPLACEMENT: {variation} -> {value}")
                
                # If XML was modified, reload the document
                if xml_content != original_xml:
                    #current_app.logger.info("🔄 XML was modified, reloading document...")
                    # Create a temporary file with the modified XML
                    import tempfile
                    import os
                    
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        tmp_path = tmp.name
                    
                    # Save the current document
                    doc.save(tmp_path)
                    
                    # Reload the document to ensure changes are applied
                    doc = Document(tmp_path)
                    
                    # Clean up temp file
                    os.unlink(tmp_path)
                    
                    #current_app.logger.info("✅ Document reloaded with XML modifications")
                else:
                    current_app.logger.info("ℹ️ No XML modifications needed")
                    
            except Exception as e:
                current_app.logger.error(f"❌ Error in nuclear XML manipulation: {e}")
                import traceback
                current_app.logger.error(f"❌ Traceback: {traceback.format_exc()}")
            
            #current_app.logger.info("✅ NUCLEAR OPTION COMPLETED")
            
            # FINAL VERIFICATION: Search for any remaining placeholders
            #current_app.logger.info("🔍 FINAL VERIFICATION: Checking for any remaining placeholders...")
            
            def search_for_remaining_placeholders(element, path=""):
                """Recursively search for any remaining placeholders"""
                if hasattr(element, 'text') and element.text:
                    text = element.text
                    
                    # Check for any remaining ${} placeholders
                    dollar_matches = re.findall(r"\$\{(.*?)\}", text)
                    for match in dollar_matches:
                        key_lower = match.lower().strip()
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ REMAINING PLACEHOLDER FOUND: ${{{match}}} in {path}")
                            current_app.logger.error(f"❌ Available data for {key_lower}: {flat_data_map.get(key_lower, 'NOT FOUND')}")
                    
                    # Check for any remaining <> placeholders
                    angle_matches = re.findall(r"<(.*?)>", text)
                    for match in angle_matches:
                        key_lower = match.lower().strip()
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ REMAINING PLACEHOLDER FOUND: <{match}> in {path}")
                            current_app.logger.error(f"❌ Available data for {key_lower}: {flat_data_map.get(key_lower, 'NOT FOUND')}")
                
                # Recursively check child elements
                for i, child in enumerate(element):
                    child_path = f"{path}.{i}" if path else str(i)
                    search_for_remaining_placeholders(child, child_path)
            
            # Search the entire document
            search_for_remaining_placeholders(doc.element, "document")
            
            #current_app.logger.info("✅ FINAL VERIFICATION COMPLETED")


                            
        # Data mapping completed silently

        def generate_chart(data_dict, chart_tag):
            import plotly.graph_objects as go
            import matplotlib.pyplot as plt
            from openpyxl.utils import column_index_from_string
            import numpy as np
            import os
            import tempfile
            import json
            import re
            import warnings
            
            # Suppress Matplotlib warnings
            warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')

            try:
                chart_tag_lower = chart_tag.lower()
                raw_chart_attr = chart_attr_map.get(chart_tag_lower, "{}")
                chart_config = json.loads(re.sub(r'//.*?\n|/\*.*?\*/', '', raw_chart_attr, flags=re.DOTALL))

                chart_meta = chart_config.get("chart_meta", {})
                series_meta = chart_config.get("series", {})
                chart_type = chart_type_map.get(chart_tag_lower, "").lower().strip()
                title = chart_meta.get("chart_title", chart_tag)

                # --- Comprehensive attribute detection logging ---
                #current_app.logger.info(f"🔍 COMPREHENSIVE CHART ATTRIBUTE DETECTION STARTED")
                #current_app.logger.info(f"📊 Chart Type: {chart_type}")
                #current_app.logger.info(f"📋 Chart Meta Keys Found: {list(chart_meta.keys())}")
                #current_app.logger.info(f"📋 Chart Config Keys Found: {list(chart_config.keys())}")
                #current_app.logger.info(f"📋 Top-level Keys Found: {list(data_dict.keys())}")
                
                # Define all possible chart attributes
                all_possible_attributes = [
                    "chart_title", "font_size", "font_color", "font_family", "figsize", 
                    "chart_background", "plot_background", "legend", "legend_position", "legend_font_size",
                    "primary_y_label", "secondary_y_label", "x_label", "y_axis_min_max", 
                    "secondary_y_axis_format", "secondary_y_axis_min_max", "x_axis_label_distance", 
                    "y_axis_label_distance", "axis_tick_format", "axis_tick_font_size",
                    "data_labels", "data_label_format", "data_label_font_size", "data_label_color",
                    "show_gridlines", "gridline_color", "gridline_style", "margin", "bar_width", 
                    "orientation", "bar_border_color", "bar_border_width", "barmode", "line_width", 
                    "marker_size", "line_style", "fill_opacity", "hole", "startangle", "pull", 
                    "bins", "box_points", "violin_points", "bubble_size", "waterfall_measure", 
                    "funnel_measure", "sunburst_path", "treemap_path", "sankey_source", 
                    "sankey_target", "sankey_value", "table_header", "indicator_mode", 
                    "indicator_delta", "indicator_gauge", "3d_projection", "z_values", 
                    "lat", "lon", "locations", "open_values", "high_values", "low_values", 
                    "close_values", "annotations"
                ]
                
                # Check which attributes are missing (not in chart_meta OR chart_config OR top-level)
                missing_attributes = []
                for attr in all_possible_attributes:
                    if (attr not in chart_meta and 
                        attr not in chart_config and 
                        attr not in data_dict):
                        missing_attributes.append(attr)
                
                # Only log missing attributes if there are any
                if missing_attributes:
                            current_app.logger.warning(f"❌ MISSING ATTRIBUTES for {chart_tag} (Not being read by system):")
                            for attr in missing_attributes:
                                current_app.logger.warning(f"   ❌ {attr}")
                            current_app.logger.warning(f"📊 Total missing: {len(missing_attributes)} attributes")
                else:
                    current_app.logger.info(f"✅ All possible attributes found for {chart_tag}")
                
                    current_app.logger.info(f"🔍 COMPREHENSIVE CHART ATTRIBUTE DETECTION COMPLETED")

                # --- Define chart type mappings for Matplotlib ---
                chart_type_mapping_mpl = {
                    # Bar charts
                    "bar": "bar",
                    "column": "bar", 
                    "stacked_column": "bar",
                    "horizontal_bar": "barh",
                    
                    # Line charts
                    "line": "plot",
                    "scatter": "scatter",
                    "scatter_line": "plot",
                    
                    # Area charts
                    "area": "fill_between",
                    "filled_area": "fill_between",
                    
                    # Statistical charts
                    "histogram": "hist",
                    "box": "boxplot",
                    "violin": "violinplot",
                    
                    # Other charts
                    "bubble": "scatter",
                    "heatmap": "imshow",
                    "contour": "contour",
                    "waterfall": "bar",
                    "funnel": "bar",
                    "sunburst": "pie",
                    "treemap": "bar",
                    "icicle": "bar",
                    "sankey": "bar",
                    "table": "table",
                    "indicator": "bar"
                }

                # --- Extract custom fields from chart_config ---
                bar_colors = chart_config.get("bar_colors")
                bar_width = data_dict.get("bar_width") or chart_config.get("bar_width") or chart_meta.get("bar_width")
                orientation = data_dict.get("orientation") or chart_config.get("orientation") or chart_meta.get("orientation")
                bar_border_color = data_dict.get("bar_border_color") or chart_config.get("bar_border_color") or chart_meta.get("bar_border_color")
                bar_border_width = data_dict.get("bar_border_width") or chart_config.get("bar_border_width") or chart_meta.get("bar_border_width")
                font_family = data_dict.get("font_family") or chart_config.get("font_family") or chart_meta.get("font_family")
                font_size = data_dict.get("font_size") or chart_config.get("font_size") or chart_meta.get("font_size")
                font_color = data_dict.get("font_color") or chart_config.get("font_color") or chart_meta.get("font_color")
                legend_position = data_dict.get("legend_position") or chart_config.get("legend_position") or chart_meta.get("legend_position")
                legend_font_size = data_dict.get("legend_font_size") or chart_config.get("legend_font_size") or chart_meta.get("legend_font_size")
                show_gridlines = data_dict.get("show_gridlines") if "show_gridlines" in data_dict else (chart_config.get("show_gridlines") if "show_gridlines" in chart_config else chart_meta.get("show_gridlines"))
                # Ensure show_gridlines is a boolean
                if isinstance(show_gridlines, str):
                    show_gridlines = show_gridlines.strip().lower() == "true"
                elif show_gridlines is None:
                    show_gridlines = True  # Default to showing gridlines if not specified
                gridline_color = data_dict.get("gridline_color") or chart_config.get("gridline_color") or chart_meta.get("gridline_color")
                gridline_style = data_dict.get("gridline_style") or chart_config.get("gridline_style") or chart_meta.get("gridline_style")
                chart_background = data_dict.get("chart_background") or chart_config.get("chart_background") or chart_meta.get("chart_background")
                plot_background = data_dict.get("plot_background") or chart_config.get("plot_background") or chart_meta.get("plot_background")
                data_label_format = data_dict.get("data_label_format") or chart_config.get("data_label_format") or chart_meta.get("data_label_format")
                data_label_font_size = data_dict.get("data_label_font_size") or chart_config.get("data_label_font_size") or chart_meta.get("data_label_font_size")
                data_label_color = data_dict.get("data_label_color") or chart_config.get("data_label_color") or chart_meta.get("data_label_color")
                axis_tick_format = data_dict.get("axis_tick_format") or chart_config.get("axis_tick_format") or chart_meta.get("axis_tick_format")
                y_axis_min_max = data_dict.get("y_axis_min_max") or chart_config.get("y_axis_min_max") or chart_meta.get("y_axis_min_max")
                # current_app.logger.debug(f"Y-axis min/max from config: {y_axis_min_max}")
                secondary_y_axis_format = data_dict.get("secondary_y_axis_format") or chart_config.get("secondary_y_axis_format") or chart_meta.get("secondary_y_axis_format")
                secondary_y_axis_min_max = data_dict.get("secondary_y_axis_min_max") or chart_config.get("secondary_y_axis_min_max") or chart_meta.get("secondary_y_axis_min_max")
                disable_secondary_y = data_dict.get("disable_secondary_y") or chart_config.get("disable_secondary_y") or chart_meta.get("disable_secondary_y", False)
                # current_app.logger.info(f"🔧 disable_secondary_y setting: {disable_secondary_y}")
                sort_order = data_dict.get("sort_order") or chart_config.get("sort_order") or chart_meta.get("sort_order")
                data_grouping = data_dict.get("data_grouping") or chart_config.get("data_grouping") or chart_meta.get("data_grouping")
                annotations = data_dict.get("annotations", []) or chart_config.get("annotations", []) or chart_meta.get("annotations", [])
                axis_tick_font_size = data_dict.get("axis_tick_font_size") or chart_config.get("axis_tick_font_size") or chart_meta.get("axis_tick_font_size")
                
                # --- Extract tick mark control settings ---
                show_x_ticks = data_dict.get("show_x_ticks") if "show_x_ticks" in data_dict else (chart_config.get("show_x_ticks") if "show_x_ticks" in chart_config else chart_meta.get("show_x_ticks"))
                show_y_ticks = data_dict.get("show_y_ticks") if "show_y_ticks" in data_dict else (chart_config.get("show_y_ticks") if "show_y_ticks" in chart_config else chart_meta.get("show_y_ticks"))
                # Ensure tick settings are boolean
                if isinstance(show_x_ticks, str):
                    show_x_ticks = show_x_ticks.strip().lower() == "true"
                elif show_x_ticks is None:
                    show_x_ticks = True  # Default to showing ticks if not specified
                if isinstance(show_y_ticks, str):
                    show_y_ticks = show_y_ticks.strip().lower() == "true"
                elif show_y_ticks is None:
                    show_y_ticks = True  # Default to showing ticks if not specified
                
                # --- Extract margin settings ---
                margin = data_dict.get("margin") or chart_config.get("margin") or chart_meta.get("margin")
                x_axis_label_distance = data_dict.get("x_axis_label_distance") or chart_config.get("x_axis_label_distance") or chart_meta.get("x_axis_label_distance")
                y_axis_label_distance = data_dict.get("y_axis_label_distance") or chart_config.get("y_axis_label_distance") or chart_meta.get("y_axis_label_distance")
                axis_tick_distance = data_dict.get("axis_tick_distance") or chart_config.get("axis_tick_distance") or chart_meta.get("axis_tick_distance")
                figsize = data_dict.get("figsize") or chart_config.get("figsize") or chart_meta.get("figsize")
                
                # --- Extract additional missing attributes ---
                legend = data_dict.get("legend") or chart_config.get("legend") or chart_meta.get("legend")
                data_labels = data_dict.get("data_labels") or chart_config.get("data_labels") or chart_meta.get("data_labels")
                line_width = data_dict.get("line_width") or chart_config.get("line_width") or chart_meta.get("line_width")
                marker_size = data_dict.get("marker_size") or chart_config.get("marker_size") or chart_meta.get("marker_size")
                line_style = data_dict.get("line_style") or chart_config.get("line_style") or chart_meta.get("line_style")
                fill_opacity = data_dict.get("fill_opacity") or chart_config.get("fill_opacity") or chart_meta.get("fill_opacity")
                hole = data_dict.get("hole") or chart_config.get("hole") or chart_meta.get("hole")
                startangle = data_dict.get("startangle") or chart_config.get("startangle") or chart_meta.get("startangle")
                pull = data_dict.get("pull") or chart_config.get("pull") or chart_meta.get("pull")
                barmode = data_dict.get("barmode") or chart_config.get("barmode") or chart_meta.get("barmode")

                # --- Excel range extraction helpers ---
                def extract_excel_range(sheet, cell_range):
                    # cell_range: e.g., 'E23:E29' or 'AA20:AA23'
                    from openpyxl.utils import range_boundaries
                    min_col, min_row, max_col, max_row = range_boundaries(cell_range)
                    values = []
                    for row in sheet.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                        for cell in row:
                            values.append(cell.value)
                    return values

                # --- Robust recursive Excel cell range extraction for all chart types and fields ---
                def extract_cell_ranges(obj, sheet):
                    """Recursively walk through nested dicts/lists and extract cell ranges from strings"""
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            if isinstance(v, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", v):
                                try:
                                    extracted = extract_excel_range(sheet, v)
                                    obj[k] = extracted
                                except Exception as e:
                                    current_app.logger.warning(f"⚠️ Failed to extract {k} from {v}: {e}")
                            else:
                                extract_cell_ranges(v, sheet)
                    elif isinstance(obj, list):
                        for i, v in enumerate(obj):
                            extract_cell_ranges(v, sheet)
                
                if "source_sheet" in chart_meta:
                    wb = openpyxl.load_workbook(data_file_path, data_only=True)
                    sheet = wb[chart_meta["source_sheet"]]
                    # Extract cell ranges from both chart_meta and series_meta
                    extract_cell_ranges(chart_meta, sheet)
                    extract_cell_ranges(series_meta, sheet)
                    wb.close()
                
                # Use updated values from series_meta after extraction
                series_data = series_meta.get("data", [])
                if not series_data and "series" in series_meta:
                    # Handle case where data is directly in series object
                    series_data = series_meta.get("series", [])
                
                # Extract x_values from the correct location
                x_values = series_meta.get("x_axis", [])
                if not x_values and "series" in series_meta:
                    # Try to get x_axis from the series object
                    x_values = series_meta.get("series", {}).get("x_axis", [])
                
                # If still no x_values, try to get from the first series data item
                if not x_values and series_data and len(series_data) > 0:
                    # Check if x_axis is in the first series
                    first_series = series_data[0]
                    if isinstance(first_series, dict) and "x_axis" in first_series:
                        x_values = first_series["x_axis"]
                    # Also check if there's a separate x_axis in the series object
                    elif "series" in series_meta and isinstance(series_meta["series"], dict):
                        x_values = series_meta["series"].get("x_axis", [])
                
                # current_app.logger.info(f"🔍 Extracted x_values: {x_values}")
                
                # Ensure x_values is always defined to prevent "cannot access local variable" error
                if not x_values:
                    x_values = []
                    current_app.logger.warning(f"⚠️ No x_values found, using empty list as fallback")
                
                colors = series_meta.get("colors", [])
                
                # --- SERIES ATTRIBUTE DETECTION LOGGING ---
                #current_app.logger.info(f"🔍 SERIES ATTRIBUTE DETECTION STARTED")
                #current_app.logger.info(f"📊 Number of series: {len(series_data)}")
                #current_app.logger.info(f"📋 Series meta keys: {list(series_data.keys()) if isinstance(series_data, dict) else 'N/A'}")
                #current_app.logger.info(f"📋 Series meta structure: {series_meta}")
                #current_app.logger.info(f"📋 Series data: {series_data}")
                #current_app.logger.info(f"📋 X values extracted: {x_values}")
                
                # Define all possible series attributes
                all_possible_series_attributes = [
                    "marker", "opacity", "textposition", "orientation", "width", "fill", 
                    "fillcolor", "hole", "pull", "mode", "line", "nbinsx", "boxpoints", 
                    "jitter", "sizeref", "sizemin", "symbol", "measure", "connector", "textinfo"
                ]
                
                for i, series in enumerate(series_data):
                    series_name = series.get("name", f"Series {i+1}")
                    series_type = series.get("type", "unknown")
                    #current_app.logger.info(f"📈 SERIES {i+1}: {series_name}")
                    #current_app.logger.info(f"   Type: {series_type}")
                    
                    # Check which series attributes are missing
                    missing_series_attributes = []
                    for attr in all_possible_series_attributes:
                        if attr not in series:
                            missing_series_attributes.append(attr)
                    
                    # Only log missing attributes if there are any
                    if missing_series_attributes:
                        current_app.logger.warning(f"   ❌ MISSING SERIES ATTRIBUTES for {series_name}:")
                        for attr in missing_series_attributes:
                            current_app.logger.warning(f"      ❌ {attr}")
                        current_app.logger.warning(f"   📊 Total missing: {len(missing_series_attributes)} attributes")
                    else:
                        current_app.logger.info(f"   ✅ All possible series attributes found for {series_name}")
                
                #current_app.logger.info(f"🔍 SERIES ATTRIBUTE DETECTION COMPLETED")

                # --- Plotly interactive chart generation ---
                fig = go.Figure()

                # --- Bar of Pie chart special handling ---
                if chart_type in ["bar of pie", "bar_of_pie"]:
                    # Extract cell ranges for other_labels and other_values if they are cell ranges
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    
                    # If other_labels and other_values are cell ranges, extract them
                    if isinstance(other_labels, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", other_labels):
                        try:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta.get("source_sheet", "sample")]
                            other_labels = extract_excel_range(sheet, other_labels)
                            wb.close()
                            current_app.logger.debug(f"Extracted other_labels from {chart_meta.get('other_labels')}: {other_labels}")
                        except Exception as e:
                            current_app.logger.warning(f"Failed to extract other_labels from {other_labels}: {e}")
                    
                    if isinstance(other_values, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", other_values):
                        try:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta.get("source_sheet", "sample")]
                            other_values = extract_excel_range(sheet, other_values)
                            wb.close()
                            current_app.logger.debug(f"Extracted other_values from {chart_meta.get('other_values')}: {other_values}")
                        except Exception as e:
                            current_app.logger.warning(f"Failed to extract other_values from {other_values}: {e}")
                    
                    # Log the extracted data for debugging
                    #current_app.logger.info(f"🔍 Bar of Pie Chart Data:")
                    #current_app.logger.info(f"   Other Labels: {other_labels}")
                    #current_app.logger.info(f"   Other Values: {other_values}")
                    #current_app.logger.info(f"   Other Colors: {other_colors}")
                    
                    # Check for empty/null values
                    if other_labels and other_values:
                        empty_count = sum(1 for label, value in zip(other_labels, other_values) 
                                        if (label is None or str(label).strip() == "" or 
                                            value is None or str(value).strip() == "" or 
                                            str(value).strip() == "0"))
                        if empty_count > 0:
                            current_app.logger.warning(f"⚠️ Found {empty_count} empty/null values in bar chart data")
                            current_app.logger.warning(f"   This will cause missing bars in the chart")
                        
                        # Log data types and values for debugging
                        current_app.logger.info(f"🔍 Data Analysis:")
                        for i, (label, value) in enumerate(zip(other_labels, other_values)):
                            current_app.logger.info(f"   Data point {i+1}: Label='{label}' (type: {type(label).__name__}), Value={value} (type: {type(value).__name__})")
                            if isinstance(value, (int, float)):
                                current_app.logger.info(f"      Numeric value: {value}")
                            elif isinstance(value, str):
                                try:
                                    numeric_val = float(value)
                                    current_app.logger.info(f"      String converted to numeric: {numeric_val}")
                                except ValueError:
                                    current_app.logger.warning(f"      Non-numeric string value: '{value}'")
                    
                    value_format = chart_meta.get("value_format", "")
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    
                    # Chart data prepared
                    fig = create_bar_of_pie_chart(
                        labels=labels,
                        values=values,
                        other_labels=other_labels,
                        other_values=other_values,
                        colors=colors,
                        other_colors=other_colors if other_colors else colors,
                        title=title,
                        value_format=value_format
                    )
                
                def extract_values_from_range(cell_range):
                    start_cell, end_cell = cell_range.split(":")
                    start_col, start_row = re.match(r"([A-Z]+)(\d+)", start_cell).groups()
                    end_col, end_row = re.match(r"([A-Z]+)(\d+)", end_cell).groups()

                    start_col_idx = column_index_from_string(start_col) - 1
                    end_col_idx = column_index_from_string(end_col) - 1
                    start_row_idx = int(start_row) - 1  # Fixed: pandas is 0-indexed, Excel is 1-indexed
                    end_row_idx = int(end_row) - 1      # Fixed: pandas is 0-indexed, Excel is 1-indexed

                    if start_col_idx == end_col_idx:
                        return df.iloc[start_row_idx:end_row_idx + 1, start_col_idx].tolist()
                    else:
                        return df.iloc[start_row_idx:end_row_idx + 1, start_col_idx:end_col_idx + 1].values.flatten().tolist()

                # --- Data grouping and sorting logic ---
                # If data_grouping is present, filter x/y values to only those groups
                def group_and_sort(x_vals, y_vals, group_names=None, sort_order=None):
                    # Grouping: only keep x/y where x in group_names (if group_names provided)
                    if group_names:
                        filtered = [(x, y) for x, y in zip(x_vals, y_vals) if x in group_names]
                        if filtered:
                            x_vals, y_vals = zip(*filtered)
                        else:
                            x_vals, y_vals = [], []
                    # Sorting: sort by y-value
                    if sort_order == "ascending":
                        sorted_pairs = sorted(zip(x_vals, y_vals), key=lambda pair: pair[1])
                        if sorted_pairs:
                            x_vals, y_vals = zip(*sorted_pairs)
                    elif sort_order == "descending":
                        sorted_pairs = sorted(zip(x_vals, y_vals), key=lambda pair: pair[1], reverse=True)
                        if sorted_pairs:
                            x_vals, y_vals = zip(*sorted_pairs)
                    return list(x_vals), list(y_vals)

                # Special handling for pie charts (single trace)
                if chart_type == "pie" and len(series_data) == 1:
                    series = series_data[0]
                    label = series.get("name", "Pie Chart")
                    labels = series.get("labels", x_values)
                    values = series.get("values", [])
                    color = series.get("marker", {}).get("color") if "marker" in series else colors
                    
                    # Apply value format if specified
                    value_format = chart_meta.get("value_format", "")
                    if value_format:
                        # Convert values to percentage format for display
                        try:
                            formatted_values = [f"{float(v):{value_format}}" if v is not None else "0" for v in values]
                        except:
                            formatted_values = values
                    else:
                        formatted_values = values
                    
                    # Check if this is an expanded pie chart (pie with one segment expanded to column)
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment:
                        # Use the helper function for expanded pie charts
                        fig = create_expanded_pie_chart(
                            labels=labels,
                            values=values,
                            colors=color if isinstance(color, list) else [color],
                            expanded_segment=expanded_segment,
                            title=title,
                            value_format=value_format
                        )
                        
                    else:
                        # Regular pie chart
                        pie_kwargs = {
                            "labels": labels,
                            "values": values,
                            "name": label,
                            "textinfo": "label+percent+value" if chart_meta.get("data_labels", True) else "none",
                            "textposition": "outside",
                            "hole": 0.0  # Solid pie chart
                        }
                        
                        # Pie colors
                        if color:
                            pie_kwargs["marker"] = dict(colors=color) if isinstance(color, list) else dict(colors=[color])
                        
                        # Add pull effect for specific segments if needed
                        if "pull" in chart_meta:
                            pie_kwargs["pull"] = chart_meta["pull"]
                        
                        fig.add_trace(go.Pie(**pie_kwargs,
                            hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}{str(value_format) if value_format else ''}<extra></extra>"
                        ))
                
                # Handle stacked column, area, and other multi-series charts
                else:
                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors):
                            color = colors[i]

                        y_vals = series.get("values")
                        value_range = series.get("value_range")
                        if value_range:
                            # Check if value_range is already extracted (list) or still a string
                            if isinstance(value_range, list):
                                y_vals = value_range
                            else:
                                y_vals = extract_values_from_range(value_range)

                        # --- Apply grouping and sorting ---
                        x_vals = x_values
                        if y_vals is not None and x_vals is not None:
                            x_vals, y_vals = group_and_sort(x_vals, y_vals, data_grouping, sort_order)

                        trace_kwargs = {
                            "x": x_vals,
                            "y": y_vals,
                            "name": label,
                        }
                        
                        # Handle colors
                        if color:
                            if isinstance(color, list) and series_type == "bar":
                                trace_kwargs["marker"] = dict(color=color)
                            elif isinstance(color, str):
                                trace_kwargs["marker_color"] = color
                        
                        # Handle bar-specific properties
                        if bar_width and series_type == "bar":
                            trace_kwargs["width"] = bar_width
                        if orientation and series_type == "bar":
                            trace_kwargs["orientation"] = orientation[0].lower() if isinstance(orientation, str) else orientation
                        if bar_border_color and series_type == "bar":
                            trace_kwargs["marker_line_color"] = bar_border_color
                        if bar_border_width and series_type == "bar":
                            trace_kwargs["marker_line_width"] = bar_border_width
                        
                        # Handle line-specific properties
                        if line_width and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["line_width"] = line_width
                        if marker_size and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["marker_size"] = marker_size
                        if line_style and series_type in ["line", "scatter", "scatter_line"]:
                            # Map line styles to Plotly format
                            line_style_map = {
                                "solid": "solid",
                                "dashed": "dash",
                                "dotted": "dot",
                                "dashdot": "dashdot"
                            }
                            trace_kwargs["line_dash"] = line_style_map.get(line_style, "solid")
                        
                        # Handle opacity
                        if fill_opacity:
                            trace_kwargs["opacity"] = fill_opacity

                        # Add traces based on chart type - REPLACE THE RESTRICTIVE IF/ELIF BLOCKS
                        # Generic chart type handling for Plotly
                        chart_type_mapping = {
                            # Bar charts
                            "bar": go.Bar,
                            "column": go.Bar,
                            "stacked_column": go.Bar,
                            "horizontal_bar": go.Bar,
                            
                            # Line charts
                            "line": go.Scatter,
                            "scatter": go.Scatter,
                            "scatter_line": go.Scatter,
                            
                            # Area charts
                            "area": go.Scatter,
                            "filled_area": go.Scatter,
                            
                            # Pie charts
                            "pie": go.Pie,
                            "donut": go.Pie,
                            
                            # 3D charts
                            "scatter3d": go.Scatter3d,
                            "surface": go.Surface,
                            "mesh3d": go.Mesh3d,
                            
                            # Statistical charts
                            "histogram": go.Histogram,
                            "box": go.Box,
                            "violin": go.Violin,
                            
                            # Financial charts
                            "candlestick": go.Candlestick,
                            "ohlc": go.Ohlc,
                            
                            # Geographic charts
                            "scattergeo": go.Scattergeo,
                            "choropleth": go.Choropleth,
                            
                            # Other charts
                            "bubble": go.Scatter,
                            "heatmap": go.Heatmap,
                            "contour": go.Contour,
                            "waterfall": go.Waterfall,
                            "funnel": go.Funnel,
                            "sunburst": go.Sunburst,
                            "treemap": go.Treemap,
                            "icicle": go.Icicle,
                            "sankey": go.Sankey,
                            "table": go.Table,
                            "indicator": go.Indicator
                        }
                        
                        # Get the appropriate Plotly chart class
                        plotly_chart_class = chart_type_mapping.get(series_type)
                        
                        if plotly_chart_class:
                            # Prepare trace arguments based on chart type
                            if series_type in ["bar", "column", "stacked_column", "horizontal_bar"]:
                                # Bar chart specific settings
                                # REMOVE: if chart_type == "stacked_column": trace_kwargs["barmode"] = "stack"
                                if orientation and orientation.lower() == "horizontal":
                                    trace_kwargs["orientation"] = "h"
                                    # Swap x and y for horizontal bars
                                    trace_kwargs["x"], trace_kwargs["y"] = trace_kwargs["y"], trace_kwargs["x"]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["line", "scatter", "scatter_line"]:
                                # Line/Scatter chart specific settings
                                mode = "lines+markers" if series_type == "scatter_line" else "markers" if series_type == "scatter" else "lines"
                                trace_kwargs["mode"] = mode
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["area", "filled_area"]:
                                # Area chart specific settings
                                trace_kwargs["mode"] = "lines"
                                trace_kwargs["fill"] = "tozeroy"
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type == "pie":
                                # Pie chart specific settings
                                pie_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": "label+percent+value" if data_labels else "none",
                                    "textposition": "outside",
                                    "hole": hole if hole is not None else (0.4 if series_type == "donut" else 0.0)
                                }
                                
                                # Add pie chart specific attributes
                                if startangle is not None:
                                    pie_kwargs["rotation"] = startangle
                                if pull is not None:
                                    # pull can be a single value or a list
                                    if isinstance(pull, (int, float)):
                                        pie_kwargs["pull"] = [pull] * len(y_vals)
                                    elif isinstance(pull, list):
                                        pie_kwargs["pull"] = pull
                                
                                if color:
                                    pie_kwargs["marker"] = dict(colors=color) if isinstance(color, list) else dict(colors=[color])
                                
                                fig.add_trace(plotly_chart_class(**pie_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            elif series_type in ["scatter3d", "surface", "mesh3d"]:
                                # 3D chart specific settings
                                if "z" not in trace_kwargs and len(y_vals) > 0:
                                    # Create a simple z-axis if not provided
                                    trace_kwargs["z"] = [i for i in range(len(y_vals))]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Z: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type in ["histogram", "box", "violin"]:
                                # Statistical chart specific settings
                                if series_type == "histogram":
                                    trace_kwargs["x"] = y_vals  # Histogram uses x for values
                                    del trace_kwargs["y"]
                                elif series_type in ["box", "violin"]:
                                    trace_kwargs["y"] = y_vals
                                    trace_kwargs["x"] = [label] * len(y_vals) if len(y_vals) > 0 else [label]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y if series_type in ['box', 'violin'] else 'x'}}<extra></extra>"
                                ))
                                
                            elif series_type in ["bubble"]:
                                # Bubble chart specific settings
                                sizes = series.get("size", [20] * len(y_vals))
                                trace_kwargs["mode"] = "markers"
                                trace_kwargs["marker"] = {"size": sizes}
                                if color:
                                    trace_kwargs["marker"]["color"] = color
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Size: %{{marker.size}}<extra></extra>"
                                ))
                                
                            elif series_type == "heatmap":
                                # Heatmap specific settings
                                heatmap_kwargs = {
                                    "z": series.get("z", y_vals),  # Use z data if provided, otherwise y_vals
                                    "x": series.get("x", x_vals),  # Use x data if provided, otherwise x_vals
                                    "y": series.get("y", [label]),  # Use y data if provided, otherwise label
                                    "name": label,
                                    "colorscale": series.get("colorscale", "Viridis"),
                                    "showscale": series.get("showscale", True)
                                }
                                
                                # Handle text data safely (convert to strings to avoid concatenation errors)
                                if "text" in series:
                                    try:
                                        heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                                    except:
                                        current_app.logger.warning(f"⚠️ Could not process heatmap text data for {label}")
                                
                                # Handle colorbar settings
                                if "colorbar" in series:
                                    heatmap_kwargs["colorbar"] = series["colorbar"]
                                
                                # Handle zmin/zmax
                                if "zmin" in series:
                                    heatmap_kwargs["zmin"] = series["zmin"]
                                if "zmax" in series:
                                    heatmap_kwargs["zmax"] = series["zmax"]
                                
                                fig.add_trace(go.Heatmap(**heatmap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                                ))
                                
                            else:
                                # Generic handling for other chart types
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y}}<extra></extra>"
                                ))
                        else:
                            # Fallback to scatter if chart type not recognized
                            current_app.logger.warning(f"⚠️ Unknown chart type '{series_type}', falling back to scatter")
                            fig.add_trace(go.Scatter(**trace_kwargs,
                                mode='markers',
                                hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                            ))

                # --- Layout updates ---
                layout_updates = {}
                
                # Title and axis labels
                layout_updates["title"] = title
                
                # Handle axis labels based on chart type
                if chart_type != "pie":
                    layout_updates["xaxis_title"] = chart_meta.get("x_label", chart_config.get("x_axis_title", "X"))
                    layout_updates["yaxis_title"] = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Y"))
                
                # Font
                if font_family or font_size or font_color:
                    layout_updates["font"] = {}
                    if font_family:
                        layout_updates["font"]["family"] = font_family
                    if font_size:
                        layout_updates["font"]["size"] = font_size
                    if font_color:
                        layout_updates["font"]["color"] = font_color
                
                # Backgrounds
                if chart_background:
                    layout_updates["paper_bgcolor"] = chart_background
                if plot_background:
                    layout_updates["plot_bgcolor"] = plot_background
                
                # Legend configuration
                show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                # current_app.logger.debug(f"Legend setting: {chart_meta.get('legend')}")
                # current_app.logger.debug(f"Showlegend setting: {chart_meta.get('showlegend')}")
                # current_app.logger.debug(f"Final show_legend: {show_legend}")
                
                if show_legend:
                    # current_app.logger.debug("Configuring legend for Plotly")
                    if legend_position:
                        # Map 'top', 'bottom', 'left', 'right' to valid Plotly legend positions
                        pos_map = {
                            "top":    dict(x=0.5, y=1.1, xanchor="center", yanchor="top"),
                            "bottom": dict(x=0.5, y=-0.2, xanchor="center", yanchor="bottom"),
                            "left":   dict(x=-0.2, y=0.5, xanchor="left", yanchor="middle"),
                            "right":  dict(x=1.1, y=0.5, xanchor="right", yanchor="middle"),
                        }
                        if legend_position in pos_map:
                            layout_updates["legend"] = pos_map[legend_position]
                            # current_app.logger.debug(f"Set legend position: {legend_position}")
                    if legend_font_size:
                        layout_updates.setdefault("legend", {})["font"] = {"size": legend_font_size}
                        # current_app.logger.debug(f"Set legend font size: {legend_font_size}")
                else:
                    layout_updates["showlegend"] = False
                    # current_app.logger.debug("Legend disabled for Plotly")
                
                # Bar mode for stacked charts
                if barmode:
                    layout_updates["barmode"] = barmode
                elif chart_type == "stacked_column":
                    layout_updates["barmode"] = "stack"
                
                # Axis min/max and tick format (only for non-pie charts)
                if chart_type != "pie":
                    if y_axis_min_max:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        # Handle "auto" value for y-axis min/max
                        if y_axis_min_max == "auto":
                            # Don't set range for auto - let Plotly auto-scale
                            pass
                        else:
                            layout_updates["yaxis"]["range"] = y_axis_min_max
                            # Also set autorange to false to ensure the range is respected
                            layout_updates["yaxis"]["autorange"] = False
                            # Force the range to be applied
                            layout_updates["yaxis"]["fixedrange"] = False
                            # Ensure the range is properly set
                            # current_app.logger.debug(f"Setting Y-axis range to: {y_axis_min_max}")
                    if axis_tick_format:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["yaxis"]["tickformat"] = axis_tick_format
                        # Also apply to secondary y-axis if it's a currency format
                        if "$" in axis_tick_format:
                            layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                            layout_updates["yaxis2"]["tickformat"] = axis_tick_format
                    
                    # Secondary y-axis formatting
                    if secondary_y_axis_format:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        layout_updates["yaxis2"]["tickformat"] = secondary_y_axis_format
                    if secondary_y_axis_min_max:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        # Handle "auto" value for secondary y-axis min/max
                        if secondary_y_axis_min_max == "auto":
                            # Don't set range for auto - let Matplotlib auto-scale
                            pass
                        elif isinstance(secondary_y_axis_min_max, list) and len(secondary_y_axis_min_max) == 2:
                            layout_updates["yaxis2"]["range"] = secondary_y_axis_min_max
                        else:
                            current_app.logger.warning(f"Invalid secondary_y_axis_min_max format: {secondary_y_axis_min_max}")
                    
                    # Axis tick font size
                    if axis_tick_font_size:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["tickfont"] = {"size": axis_tick_font_size}
                        layout_updates["yaxis"]["tickfont"] = {"size": axis_tick_font_size}
                    # Gridlines
                    if show_gridlines is not None:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["showgrid"] = bool(show_gridlines)
                        layout_updates["yaxis"]["showgrid"] = bool(show_gridlines)
                    if gridline_color:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["gridcolor"] = gridline_color
                        layout_updates["yaxis"]["gridcolor"] = gridline_color
                    if gridline_style:
                        # Map gridline styles to valid Plotly dash styles
                        dash_map = {
                            "solid": "solid", 
                            "dashed": "dash",
                            "dash": "dash",  # Map 'dash' to 'dash'
                            "dashdot": "dashdot", 
                            "dotted": "dot",
                            "dot": "dot",
                            "dotdash": "dashdot"
                        }
                        dash_style = dash_map.get(gridline_style, "solid")
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["griddash"] = dash_style
                        layout_updates["yaxis"]["griddash"] = dash_style
                
                # Data labels (for bar/line traces)
                show_data_labels = chart_meta.get("data_labels", True)
                value_format = chart_meta.get("value_format", "")
                
                # Only enable data labels if explicitly set to True AND any data label settings are provided
                if show_data_labels and (data_label_format or data_label_font_size or data_label_color):
                    show_data_labels = True
                elif not show_data_labels:
                    # If data_labels is explicitly set to False, respect that setting
                    show_data_labels = False
                
                # Debug logging for data labels
                # current_app.logger.debug(f"Original data_labels setting: {chart_meta.get('data_labels')}")
                # current_app.logger.debug(f"Final show_data_labels: {show_data_labels}")
                # current_app.logger.debug(f"Data label format: {data_label_format}")
                # current_app.logger.debug(f"Data label font size: {data_label_font_size}")
                # current_app.logger.debug(f"Data label color: {data_label_color}")
                # current_app.logger.debug(f"Value format: {value_format}")
                # current_app.logger.debug(f"Chart config keys: {list(chart_config.keys())}")
                # current_app.logger.debug(f"Chart meta keys: {list(chart_meta.keys())}")
                
                if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color):
                    # current_app.logger.debug(f"Processing {len(fig.data)} traces for data labels")
                    for i, trace in enumerate(fig.data):
                        # current_app.logger.debug(f"Trace {i}: type={trace.type}, mode={getattr(trace, 'mode', 'N/A')}")
                        # Handle both bar and line charts (line charts are scatter with mode='lines')
                        if trace.type in ['bar', 'scatter']:
                            # Use value_format from chart_meta if available, otherwise use data_label_format
                            format_to_use = value_format if value_format else data_label_format
                            
                            # Determine if this is a line chart (scatter with lines mode)
                            is_line_chart = trace.type == 'scatter' and trace.mode and 'lines' in trace.mode
                            
                            if format_to_use:
                                if trace.type == 'bar':
                                    trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                            else:
                                # If no format specified, still show data labels
                                if trace.type == 'bar':
                                    trace.update(texttemplate="%{y}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate="%{y}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate="%{y}", textposition="top center")
                            
                            # Apply font styling
                            if data_label_font_size or data_label_color:
                                trace.update(textfont={})
                                if data_label_font_size:
                                    trace.textfont["size"] = data_label_font_size
                                if data_label_color:
                                    trace.textfont["color"] = data_label_color
                
                # Annotations
                if annotations:
                    layout_updates["annotations"] = []
                    for ann in annotations:
                        ann_dict = {
                            "text": ann.get("text", ""),
                            "x": ann.get("x_value", ann.get("x")),
                            "y": ann.get("y_value", ann.get("y")),
                            "showarrow": True
                        }
                        layout_updates["annotations"].append(ann_dict)
                
                # --- Apply margin settings ---
                if margin:
                    layout_updates["margin"] = margin
                
                # Apply figure size if specified
                if figsize:
                    layout_updates["width"] = figsize[0] * 100  # Convert to pixels
                    layout_updates["height"] = figsize[1] * 100  # Convert to pixels
                    # current_app.logger.debug(f"Applied Plotly figsize: {figsize} -> width={figsize[0]*100}, height={figsize[1]*100}")
                
                # Apply axis label distances
                if chart_type != "pie":
                    if x_axis_label_distance or y_axis_label_distance or axis_tick_distance:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        
                        if x_axis_label_distance:
                            layout_updates["xaxis"]["title"] = layout_updates["xaxis"].get("title", {})
                            layout_updates["xaxis"]["title"]["standoff"] = x_axis_label_distance
                            # Also apply to tick distance for better control
                            layout_updates["xaxis"]["ticklen"] = x_axis_label_distance
                        
                        if y_axis_label_distance:
                            layout_updates["yaxis"]["title"] = layout_updates["yaxis"].get("title", {})
                            layout_updates["yaxis"]["title"]["standoff"] = y_axis_label_distance
                        
                        if axis_tick_distance:
                            layout_updates["xaxis"]["ticklen"] = axis_tick_distance
                            layout_updates["yaxis"]["ticklen"] = axis_tick_distance

                # Tick mark control for Plotly
                if show_x_ticks is not None or show_y_ticks is not None:
                    layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                    layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                    if show_x_ticks is not None:
                        layout_updates["xaxis"]["showticklabels"] = bool(show_x_ticks)
                        layout_updates["xaxis"]["ticks"] = "" if not show_x_ticks else "outside"
                    if show_y_ticks is not None:
                        layout_updates["yaxis"]["showticklabels"] = bool(show_y_ticks)
                        layout_updates["yaxis"]["ticks"] = "" if not show_y_ticks else "outside"

                fig.update_layout(**layout_updates)

                # --- Matplotlib static chart for DOCX ---
                if chart_type == "pie":
                    # Check if this is an expanded pie chart
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment and len(series_data) == 1:
                        # Create subplot for expanded pie chart
                        mpl_figsize = figsize if figsize else (15, 8)
                        fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax1.set_facecolor(plot_background)
                            ax2.set_facecolor(plot_background)
                        
                        series = series_data[0]
                        labels = series.get("labels", x_values)
                        values = series.get("values", [])
                        color = series.get("marker", {}).get("color") if "marker" in series else colors
                        
                        # Create pie chart
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', 
                                                          colors=color, startangle=90)
                        
                        # Style the text
                        for autotext in autotexts:
                            autotext.set_color('white')
                            autotext.set_fontweight('bold')
                        
                        ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                        
                        # Add legend for pie chart
                        show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        if show_legend:
                            # Initialize legend_loc for pie charts
                            legend_loc = 'best'  # default
                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right"
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                            
                            # Force legend to bottom if specified
                            if legend_position == "bottom":
                                ax1.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                            else:
                                ax1.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                        # Create bar chart for expanded segment
                        if expanded_segment in labels:
                            segment_idx = labels.index(expanded_segment)
                            segment_value = values[segment_idx]
                            segment_color = color[segment_idx] if isinstance(color, list) and segment_idx < len(color) else color
                            
                            ax2.bar([expanded_segment], [segment_value], color=segment_color, alpha=0.7)
                            ax2.set_title(f"{expanded_segment} Details", fontsize=font_size or 12, weight='bold')
                            ax2.set_ylabel("Value")
                            
                            # Add value label on bar
                            ax2.text(0, segment_value, f"{segment_value}", ha='center', va='bottom', fontweight='bold')
                        
                    else:
                        # Regular pie chart
                        mpl_figsize = figsize if figsize else (10, 8)
                        fig_mpl, ax = plt.subplots(figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax.set_facecolor(plot_background)
                        
                        if len(series_data) == 1:
                            series = series_data[0]
                            labels = series.get("labels", x_values)
                            values = series.get("values", [])
                            color = series.get("marker", {}).get("color") if "marker" in series else colors
                            
                            # Create pie chart
                            wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', 
                                                             colors=color, startangle=90)
                            
                            # Style the text
                            for autotext in autotexts:
                                autotext.set_color('white')
                                autotext.set_fontweight('bold')
                            
                            ax.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                            
                            # Add legend for regular pie chart
                            show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            if show_legend:
                                # Initialize legend_loc for pie charts
                                legend_loc = 'best'  # default
                                if legend_position:
                                    loc_map = {
                                        "top": "upper center",
                                        "bottom": "lower center", 
                                        "left": "center left",
                                        "right": "center right"
                                    }
                                    legend_loc = loc_map.get(legend_position, 'best')
                                
                                # Force legend to bottom if specified
                                if legend_position == "bottom":
                                    ax.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                else:
                                    ax.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                elif chart_type in ["bar of pie", "bar_of_pie"]:
                    # Matplotlib version of bar of pie
                    mpl_figsize = figsize if figsize else (10, 5)
                    fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200, gridspec_kw={'width_ratios': [2, 1]})
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    if not (other_labels and other_values):
                        if "other_label_range" in chart_meta and "other_value_range" in chart_meta and "source_sheet" in chart_meta:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta["source_sheet"]]
                            other_labels = extract_excel_range(sheet, chart_meta["other_label_range"])
                            other_values = extract_excel_range(sheet, chart_meta["other_value_range"])
                    # Pie chart
                    if colors:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
                    else:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')
                    ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                    # Move legend outside
                    show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                    legend_font_size = chart_meta.get("legend_font_size", 8)
                    if show_legend:
                        ax1.legend(wedges, labels, loc='center left', bbox_to_anchor=(1, 0.5), fontsize=legend_font_size)
                    
                    # Bar chart - Filter out empty/null values first
                    filtered_data = []
                    for label, value in zip(other_labels, other_values):
                        if (label is not None and str(label).strip() != "" and 
                            value is not None and str(value).strip() != "" and 
                            str(value).strip() != "0"):
                            filtered_data.append((label, value))
                    
                    if filtered_data:
                        filtered_labels, filtered_values = zip(*filtered_data)
                    else:
                        filtered_labels, filtered_values = [], []
                    
                    # Create individual bars (not stacked)
                    if filtered_labels and filtered_values:
                        # Use individual colors for each bar
                        bar_colors = []
                        for i in range(len(filtered_labels)):
                            if other_colors and i < len(other_colors):
                                bar_colors.append(other_colors[i])
                            elif colors and i < len(colors):
                                bar_colors.append(colors[i])
                            else:
                                bar_colors.append('#1f77b4')  # Default blue
                        
                        bars = ax2.bar(range(len(filtered_labels)), filtered_values, color=bar_colors, alpha=0.7)
                    else:
                        bars = []
                    
                    expanded_segment = chart_meta.get("expanded_segment", "Other")
                    ax2.set_title(f"Breakdown of '{expanded_segment}'", fontsize=font_size or 12, weight='bold')
                    # Use proper Y-axis label from configuration
                    y_axis_title = chart_meta.get("y_axis_title", "Value")
                    ax2.set_ylabel(y_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set X-axis label
                    x_axis_title = chart_meta.get("x_axis_title", "Categories")
                    ax2.set_xlabel(x_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set x-tick labels for filtered data
                    if filtered_labels:
                        ax2.set_xticks(range(len(filtered_labels)))
                        ax2.set_xticklabels(filtered_labels, rotation=0)
                    # Add data labels with proper formatting
                    value_format = chart_meta.get("value_format", ".2f")
                    data_label_font_size = chart_meta.get("data_label_font_size", 10)
                    data_label_color = chart_meta.get("data_label_color", "#000000")
                    for bar, v in zip(bars, filtered_values):
                        if value_format == "value":
                            formatted_value = str(v)
                        elif value_format == ".2f":
                            formatted_value = f"{v:.2f}"
                        elif value_format == ".1f":
                            formatted_value = f"{v:.1f}"
                        elif value_format == ".0f":
                            formatted_value = f"{v:.0f}"
                        else:
                            formatted_value = f"{v:{value_format}}"
                        ax2.text(bar.get_x() + bar.get_width()/2, v, formatted_value, ha='center', va='bottom', fontweight='bold', fontsize=data_label_font_size, color=data_label_color)

                else:
                    # Bar, line, area charts
                    mpl_figsize = figsize if figsize else (10, 6)
                    # current_app.logger.debug(f"Applied Matplotlib figsize: {mpl_figsize}")
                    fig_mpl, ax1 = plt.subplots(figsize=mpl_figsize, dpi=200)
                    ax2 = ax1.twinx()
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)

                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors):
                            color = colors[i]

                        y_vals = series.get("values")
                        value_range = series.get("value_range")
                        if value_range:
                            # Check if value_range is already extracted (list) or still a string
                            if isinstance(value_range, list):
                                y_vals = value_range
                            else:
                                y_vals = extract_values_from_range(value_range)

                        # --- Apply grouping and sorting ---
                        x_vals = x_values
                        if y_vals is not None and x_vals is not None:
                            x_vals, y_vals = group_and_sort(x_vals, y_vals, data_grouping, sort_order)

                        trace_kwargs = {
                            "x": x_vals,
                            "y": y_vals,
                            "name": label,
                        }
                        
                        # Handle colors
                        if color:
                            if isinstance(color, list) and series_type == "bar":
                                trace_kwargs["marker"] = dict(color=color)
                            elif isinstance(color, str):
                                trace_kwargs["marker_color"] = color
                        
                        # Handle bar-specific properties
                        if bar_width and series_type == "bar":
                            trace_kwargs["width"] = bar_width
                        if orientation and series_type == "bar":
                            trace_kwargs["orientation"] = orientation[0].lower() if isinstance(orientation, str) else orientation
                        if bar_border_color and series_type == "bar":
                            trace_kwargs["marker_line_color"] = bar_border_color
                        if bar_border_width and series_type == "bar":
                            trace_kwargs["marker_line_width"] = bar_border_width
                        
                        # Handle line-specific properties
                        if line_width and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["line_width"] = line_width
                        if marker_size and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["marker_size"] = marker_size
                        if line_style and series_type in ["line", "scatter", "scatter_line"]:
                            # Map line styles to Plotly format
                            line_style_map = {
                                "solid": "solid",
                                "dashed": "dash",
                                "dotted": "dot",
                                "dashdot": "dashdot"
                            }
                            trace_kwargs["line_dash"] = line_style_map.get(line_style, "solid")
                        
                        # Handle opacity
                        if fill_opacity:
                            trace_kwargs["opacity"] = fill_opacity

                        # Add traces based on chart type - REPLACE THE RESTRICTIVE IF/ELIF BLOCKS
                        # Generic chart type handling for Plotly
                        chart_type_mapping = {
                            # Bar charts
                            "bar": go.Bar,
                            "column": go.Bar,
                            "stacked_column": go.Bar,
                            "horizontal_bar": go.Bar,
                            
                            # Line charts
                            "line": go.Scatter,
                            "scatter": go.Scatter,
                            "scatter_line": go.Scatter,
                            
                            # Area charts
                            "area": go.Scatter,
                            "filled_area": go.Scatter,
                            
                            # Pie charts
                            "pie": go.Pie,
                            "donut": go.Pie,
                            
                            # 3D charts
                            "scatter3d": go.Scatter3d,
                            "surface": go.Surface,
                            "mesh3d": go.Mesh3d,
                            
                            # Statistical charts
                            "histogram": go.Histogram,
                            "box": go.Box,
                            "violin": go.Violin,
                            
                            # Financial charts
                            "candlestick": go.Candlestick,
                            "ohlc": go.Ohlc,
                            
                            # Geographic charts
                            "scattergeo": go.Scattergeo,
                            "choropleth": go.Choropleth,
                            
                            # Other charts
                            "bubble": go.Scatter,
                            "heatmap": go.Heatmap,
                            "contour": go.Contour,
                            "waterfall": go.Waterfall,
                            "funnel": go.Funnel,
                            "sunburst": go.Sunburst,
                            "treemap": go.Treemap,
                            "icicle": go.Icicle,
                            "sankey": go.Sankey,
                            "table": go.Table,
                            "indicator": go.Indicator
                        }
                        
                        # Get the appropriate Plotly chart class
                        plotly_chart_class = chart_type_mapping.get(series_type)
                        
                        if plotly_chart_class:
                            # Prepare trace arguments based on chart type
                            if series_type in ["bar", "column", "stacked_column", "horizontal_bar"]:
                                # Bar chart specific settings
                                # REMOVE: if chart_type == "stacked_column": trace_kwargs["barmode"] = "stack"
                                if orientation and orientation.lower() == "horizontal":
                                    trace_kwargs["orientation"] = "h"
                                    # Swap x and y for horizontal bars
                                    trace_kwargs["x"], trace_kwargs["y"] = trace_kwargs["y"], trace_kwargs["x"]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["line", "scatter", "scatter_line"]:
                                # Line/Scatter chart specific settings
                                mode = "lines+markers" if series_type == "scatter_line" else "markers" if series_type == "scatter" else "lines"
                                trace_kwargs["mode"] = mode
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["area", "filled_area"]:
                                # Area chart specific settings
                                trace_kwargs["mode"] = "lines"
                                trace_kwargs["fill"] = "tozeroy"
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type == "pie":
                                # Pie chart specific settings
                                pie_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": "label+percent+value" if data_labels else "none",
                                    "textposition": "outside",
                                    "hole": hole if hole is not None else (0.4 if series_type == "donut" else 0.0)
                                }
                                
                                # Add pie chart specific attributes
                                if startangle is not None:
                                    pie_kwargs["rotation"] = startangle
                                if pull is not None:
                                    # pull can be a single value or a list
                                    if isinstance(pull, (int, float)):
                                        pie_kwargs["pull"] = [pull] * len(y_vals)
                                    elif isinstance(pull, list):
                                        pie_kwargs["pull"] = pull
                                
                                if color:
                                    pie_kwargs["marker"] = dict(colors=color) if isinstance(color, list) else dict(colors=[color])
                                
                                fig.add_trace(plotly_chart_class(**pie_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            elif series_type in ["scatter3d", "surface", "mesh3d"]:
                                # 3D chart specific settings
                                if "z" not in trace_kwargs and len(y_vals) > 0:
                                    # Create a simple z-axis if not provided
                                    trace_kwargs["z"] = [i for i in range(len(y_vals))]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Z: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type in ["histogram", "box", "violin"]:
                                # Statistical chart specific settings
                                if series_type == "histogram":
                                    trace_kwargs["x"] = y_vals  # Histogram uses x for values
                                    del trace_kwargs["y"]
                                elif series_type in ["box", "violin"]:
                                    trace_kwargs["y"] = y_vals
                                    trace_kwargs["x"] = [label] * len(y_vals) if len(y_vals) > 0 else [label]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y if series_type in ['box', 'violin'] else 'x'}}<extra></extra>"
                                ))
                                
                            elif series_type in ["bubble"]:
                                # Bubble chart specific settings
                                sizes = series.get("size", [20] * len(y_vals))
                                trace_kwargs["mode"] = "markers"
                                trace_kwargs["marker"] = {"size": sizes}
                                if color:
                                    trace_kwargs["marker"]["color"] = color
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Size: %{{marker.size}}<extra></extra>"
                                ))
                                
                            elif series_type == "heatmap":
                                # Heatmap specific settings
                                heatmap_kwargs = {
                                    "z": series.get("z", y_vals),  # Use z data if provided, otherwise y_vals
                                    "x": series.get("x", x_vals),  # Use x data if provided, otherwise x_vals
                                    "y": series.get("y", [label]),  # Use y data if provided, otherwise label
                                    "name": label,
                                    "colorscale": series.get("colorscale", "Viridis"),
                                    "showscale": series.get("showscale", True)
                                }
                                
                                # Handle text data safely (convert to strings to avoid concatenation errors)
                                if "text" in series:
                                    try:
                                        heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                                    except:
                                        current_app.logger.warning(f"⚠️ Could not process heatmap text data for {label}")
                                
                                # Handle colorbar settings
                                if "colorbar" in series:
                                    heatmap_kwargs["colorbar"] = series["colorbar"]
                                
                                # Handle zmin/zmax
                                if "zmin" in series:
                                    heatmap_kwargs["zmin"] = series["zmin"]
                                if "zmax" in series:
                                    heatmap_kwargs["zmax"] = series["zmax"]
                                
                                fig.add_trace(go.Heatmap(**heatmap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                                ))
                                
                            else:
                                # Generic handling for other chart types
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y}}<extra></extra>"
                                ))
                        else:
                            # Fallback to scatter if chart type not recognized
                            current_app.logger.warning(f"⚠️ Unknown chart type '{series_type}', falling back to scatter")
                            fig.add_trace(go.Scatter(**trace_kwargs,
                                mode='markers',
                                hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                            ))

                # --- Process heatmap charts separately ---
                # Heatmaps need special handling because they have x, y, z, text data directly
                for i, series in enumerate(series_data):
                    label = series.get("name", f"Series {i+1}")
                    series_type = series.get("type", "bar").lower()
                    
                    if series_type == "heatmap":
                        # Heatmap specific settings
                        heatmap_kwargs = {
                            "z": series.get("z", []),  # Use z data directly
                            "x": series.get("x", []),  # Use x data directly
                            "y": series.get("y", []),  # Use y data directly
                            "name": label,
                            "colorscale": series.get("colorscale", "Viridis"),
                            "showscale": series.get("showscale", True)
                        }
                        
                        # Handle text data safely (convert to strings to avoid concatenation errors)
                        if "text" in series:
                            try:
                                heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                            except:
                                current_app.logger.warning(f"⚠️ Could not process heatmap text data for {label}")
                        
                        # Handle colorbar settings
                        if "colorbar" in series:
                            heatmap_kwargs["colorbar"] = series["colorbar"]
                        
                        # Handle zmin/zmax
                        if "zmin" in series:
                            heatmap_kwargs["zmin"] = series["zmin"]
                        if "zmax" in series:
                            heatmap_kwargs["zmax"] = series["zmax"]
                        
                        fig.add_trace(go.Heatmap(**heatmap_kwargs,
                            hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                        ))

                # --- Layout updates ---
                layout_updates = {}
                
                # Title and axis labels
                layout_updates["title"] = title
                
                # Handle axis labels based on chart type
                if chart_type != "pie":
                    layout_updates["xaxis_title"] = chart_meta.get("x_label", chart_config.get("x_axis_title", "X"))
                    layout_updates["yaxis_title"] = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Y"))
                
                # Font
                if font_family or font_size or font_color:
                    layout_updates["font"] = {}
                    if font_family:
                        layout_updates["font"]["family"] = font_family
                    if font_size:
                        layout_updates["font"]["size"] = font_size
                    if font_color:
                        layout_updates["font"]["color"] = font_color
                
                # Backgrounds
                if chart_background:
                    layout_updates["paper_bgcolor"] = chart_background
                if plot_background:
                    layout_updates["plot_bgcolor"] = plot_background
                
                # Legend configuration
                show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                # current_app.logger.debug(f"Legend setting: {chart_meta.get('legend')}")
                # current_app.logger.debug(f"Showlegend setting: {chart_meta.get('showlegend')}")
                # current_app.logger.debug(f"Final show_legend: {show_legend}")
                
                if show_legend:
                    # current_app.logger.debug("Configuring legend for Plotly")
                    if legend_position:
                        # Map 'top', 'bottom', 'left', 'right' to valid Plotly legend positions
                        pos_map = {
                            "top":    dict(x=0.5, y=1.1, xanchor="center", yanchor="top"),
                            "bottom": dict(x=0.5, y=-0.2, xanchor="center", yanchor="bottom"),
                            "left":   dict(x=-0.2, y=0.5, xanchor="left", yanchor="middle"),
                            "right":  dict(x=1.1, y=0.5, xanchor="right", yanchor="middle"),
                        }
                        if legend_position in pos_map:
                            layout_updates["legend"] = pos_map[legend_position]
                            # current_app.logger.debug(f"Set legend position: {legend_position}")
                    if legend_font_size:
                        layout_updates.setdefault("legend", {})["font"] = {"size": legend_font_size}
                        # current_app.logger.debug(f"Set legend font size: {legend_font_size}")
                else:
                    layout_updates["showlegend"] = False
                    # current_app.logger.debug("Legend disabled for Plotly")
                
                # Bar mode for stacked charts
                if barmode:
                    layout_updates["barmode"] = barmode
                elif chart_type == "stacked_column":
                    layout_updates["barmode"] = "stack"
                
                # Axis min/max and tick format (only for non-pie charts)
                if chart_type != "pie":
                    if y_axis_min_max:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        # Handle "auto" value for y-axis min/max
                        if y_axis_min_max == "auto":
                            # Don't set range for auto - let Plotly auto-scale
                            pass
                        else:
                            layout_updates["yaxis"]["range"] = y_axis_min_max
                            # Also set autorange to false to ensure the range is respected
                            layout_updates["yaxis"]["autorange"] = False
                            # Force the range to be applied
                            layout_updates["yaxis"]["fixedrange"] = False
                            # Ensure the range is properly set
                            # current_app.logger.debug(f"Setting Y-axis range to: {y_axis_min_max}")
                    if axis_tick_format:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["yaxis"]["tickformat"] = axis_tick_format
                        # Also apply to secondary y-axis if it's a currency format
                        if "$" in axis_tick_format:
                            layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                            layout_updates["yaxis2"]["tickformat"] = axis_tick_format
                    
                    # Secondary y-axis formatting
                    if secondary_y_axis_format:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        layout_updates["yaxis2"]["tickformat"] = secondary_y_axis_format
                    if secondary_y_axis_min_max:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        # Handle "auto" value for secondary y-axis min/max
                        if secondary_y_axis_min_max == "auto":
                            # Don't set range for auto - let Matplotlib auto-scale
                            pass
                        elif isinstance(secondary_y_axis_min_max, list) and len(secondary_y_axis_min_max) == 2:
                            layout_updates["yaxis2"]["range"] = secondary_y_axis_min_max
                        else:
                            current_app.logger.warning(f"Invalid secondary_y_axis_min_max format: {secondary_y_axis_min_max}")
                    
                    # Axis tick font size
                    if axis_tick_font_size:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["tickfont"] = {"size": axis_tick_font_size}
                        layout_updates["yaxis"]["tickfont"] = {"size": axis_tick_font_size}
                    # Gridlines
                    if show_gridlines is not None:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["showgrid"] = bool(show_gridlines)
                        layout_updates["yaxis"]["showgrid"] = bool(show_gridlines)
                    if gridline_color:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["gridcolor"] = gridline_color
                        layout_updates["yaxis"]["gridcolor"] = gridline_color
                    if gridline_style:
                        # Map gridline styles to valid Plotly dash styles
                        dash_map = {
                            "solid": "solid", 
                            "dashed": "dash",
                            "dash": "dash",  # Map 'dash' to 'dash'
                            "dashdot": "dashdot", 
                            "dotted": "dot",
                            "dot": "dot",
                            "dotdash": "dashdot"
                        }
                        dash_style = dash_map.get(gridline_style, "solid")
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["griddash"] = dash_style
                        layout_updates["yaxis"]["griddash"] = dash_style
                
                # Data labels (for bar/line traces)
                show_data_labels = chart_meta.get("data_labels", True)
                value_format = chart_meta.get("value_format", "")
                
                # Only enable data labels if explicitly set to True AND any data label settings are provided
                if show_data_labels and (data_label_format or data_label_font_size or data_label_color):
                    show_data_labels = True
                elif not show_data_labels:
                    # If data_labels is explicitly set to False, respect that setting
                    show_data_labels = False
                
                # Debug logging for data labels
                # current_app.logger.debug(f"Original data_labels setting: {chart_meta.get('data_labels')}")
                # current_app.logger.debug(f"Final show_data_labels: {show_data_labels}")
                # current_app.logger.debug(f"Data label format: {data_label_format}")
                # current_app.logger.debug(f"Data label font size: {data_label_font_size}")
                # current_app.logger.debug(f"Data label color: {data_label_color}")
                # current_app.logger.debug(f"Value format: {value_format}")
                # current_app.logger.debug(f"Chart config keys: {list(chart_config.keys())}")
                # current_app.logger.debug(f"Chart meta keys: {list(chart_meta.keys())}")
                
                if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color):
                    # current_app.logger.debug(f"Processing {len(fig.data)} traces for data labels")
                    for i, trace in enumerate(fig.data):
                        # current_app.logger.debug(f"Trace {i}: type={trace.type}, mode={getattr(trace, 'mode', 'N/A')}")
                        # Handle both bar and line charts (line charts are scatter with mode='lines')
                        if trace.type in ['bar', 'scatter']:
                            # Use value_format from chart_meta if available, otherwise use data_label_format
                            format_to_use = value_format if value_format else data_label_format
                            
                            # Determine if this is a line chart (scatter with lines mode)
                            is_line_chart = trace.type == 'scatter' and trace.mode and 'lines' in trace.mode
                            
                            if format_to_use:
                                if trace.type == 'bar':
                                    trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                            else:
                                # If no format specified, still show data labels
                                if trace.type == 'bar':
                                    trace.update(texttemplate="%{y}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate="%{y}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate="%{y}", textposition="top center")
                            
                            # Apply font styling
                            if data_label_font_size or data_label_color:
                                trace.update(textfont={})
                                if data_label_font_size:
                                    trace.textfont["size"] = data_label_font_size
                                if data_label_color:
                                    trace.textfont["color"] = data_label_color
                
                # Annotations
                if annotations:
                    layout_updates["annotations"] = []
                    for ann in annotations:
                        ann_dict = {
                            "text": ann.get("text", ""),
                            "x": ann.get("x_value", ann.get("x")),
                            "y": ann.get("y_value", ann.get("y")),
                            "showarrow": True
                        }
                        layout_updates["annotations"].append(ann_dict)
                
                # --- Apply margin settings ---
                if margin:
                    layout_updates["margin"] = margin
                
                # Apply figure size if specified
                if figsize:
                    layout_updates["width"] = figsize[0] * 100  # Convert to pixels
                    layout_updates["height"] = figsize[1] * 100  # Convert to pixels
                    # current_app.logger.debug(f"Applied Plotly figsize: {figsize} -> width={figsize[0]*100}, height={figsize[1]*100}")
                
                # Apply axis label distances
                if chart_type != "pie":
                    if x_axis_label_distance or y_axis_label_distance or axis_tick_distance:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        
                        if x_axis_label_distance:
                            layout_updates["xaxis"]["title"] = layout_updates["xaxis"].get("title", {})
                            layout_updates["xaxis"]["title"]["standoff"] = x_axis_label_distance
                            # Also apply to tick distance for better control
                            layout_updates["xaxis"]["ticklen"] = x_axis_label_distance
                        
                        if y_axis_label_distance:
                            layout_updates["yaxis"]["title"] = layout_updates["yaxis"].get("title", {})
                            layout_updates["yaxis"]["title"]["standoff"] = y_axis_label_distance
                        
                        if axis_tick_distance:
                            layout_updates["xaxis"]["ticklen"] = axis_tick_distance
                            layout_updates["yaxis"]["ticklen"] = axis_tick_distance

                # Tick mark control for Plotly
                if show_x_ticks is not None or show_y_ticks is not None:
                    layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                    layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                    if show_x_ticks is not None:
                        layout_updates["xaxis"]["showticklabels"] = bool(show_x_ticks)
                        layout_updates["xaxis"]["ticks"] = "" if not show_x_ticks else "outside"
                    if show_y_ticks is not None:
                        layout_updates["yaxis"]["showticklabels"] = bool(show_y_ticks)
                        layout_updates["yaxis"]["ticks"] = "" if not show_y_ticks else "outside"

                fig.update_layout(**layout_updates)

                # --- Matplotlib static chart for DOCX ---
                if chart_type == "pie":
                    # Check if this is an expanded pie chart
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment and len(series_data) == 1:
                        # Create subplot for expanded pie chart
                        mpl_figsize = figsize if figsize else (15, 8)
                        fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax1.set_facecolor(plot_background)
                            ax2.set_facecolor(plot_background)
                        
                        series = series_data[0]
                        labels = series.get("labels", x_values)
                        values = series.get("values", [])
                        color = series.get("marker", {}).get("color") if "marker" in series else colors
                        
                        # Create pie chart
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', 
                                                          colors=color, startangle=90)
                        
                        # Style the text
                        for autotext in autotexts:
                            autotext.set_color('white')
                            autotext.set_fontweight('bold')
                        
                        ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                        
                        # Add legend for pie chart
                        show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        if show_legend:
                            # Initialize legend_loc for pie charts
                            legend_loc = 'best'  # default
                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right"
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                            
                            # Force legend to bottom if specified
                            if legend_position == "bottom":
                                ax1.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                            else:
                                ax1.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                        # Create bar chart for expanded segment
                        if expanded_segment in labels:
                            segment_idx = labels.index(expanded_segment)
                            segment_value = values[segment_idx]
                            segment_color = color[segment_idx] if isinstance(color, list) and segment_idx < len(color) else color
                            
                            ax2.bar([expanded_segment], [segment_value], color=segment_color, alpha=0.7)
                            ax2.set_title(f"{expanded_segment} Details", fontsize=font_size or 12, weight='bold')
                            ax2.set_ylabel("Value")
                            
                            # Add value label on bar
                            ax2.text(0, segment_value, f"{segment_value}", ha='center', va='bottom', fontweight='bold')
                        
                    else:
                        # Regular pie chart
                        mpl_figsize = figsize if figsize else (10, 8)
                        fig_mpl, ax = plt.subplots(figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax.set_facecolor(plot_background)
                        
                        if len(series_data) == 1:
                            series = series_data[0]
                            labels = series.get("labels", x_values)
                            values = series.get("values", [])
                            color = series.get("marker", {}).get("color") if "marker" in series else colors
                            
                            # Create pie chart
                            wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', 
                                                             colors=color, startangle=90)
                            
                            # Style the text
                            for autotext in autotexts:
                                autotext.set_color('white')
                                autotext.set_fontweight('bold')
                            
                            ax.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                            
                            # Add legend for regular pie chart
                            show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            if show_legend:
                                # Initialize legend_loc for pie charts
                                legend_loc = 'best'  # default
                                if legend_position:
                                    loc_map = {
                                        "top": "upper center",
                                        "bottom": "lower center", 
                                        "left": "center left",
                                        "right": "center right"
                                    }
                                    legend_loc = loc_map.get(legend_position, 'best')
                                
                                # Force legend to bottom if specified
                                if legend_position == "bottom":
                                    ax.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                else:
                                    ax.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                elif chart_type in ["bar of pie", "bar_of_pie"]:
                    # Matplotlib version of bar of pie
                    mpl_figsize = figsize if figsize else (10, 5)
                    fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200, gridspec_kw={'width_ratios': [2, 1]})
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    if not (other_labels and other_values):
                        if "other_label_range" in chart_meta and "other_value_range" in chart_meta and "source_sheet" in chart_meta:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta["source_sheet"]]
                            other_labels = extract_excel_range(sheet, chart_meta["other_label_range"])
                            other_values = extract_excel_range(sheet, chart_meta["other_value_range"])
                    # Pie chart
                    if colors:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
                    else:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')
                    ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                    # Move legend outside
                    show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                    legend_font_size = chart_meta.get("legend_font_size", 8)
                    if show_legend:
                        ax1.legend(wedges, labels, loc='center left', bbox_to_anchor=(1, 0.5), fontsize=legend_font_size)
                    
                    # Bar chart - Filter out empty/null values first
                    filtered_data = []
                    for label, value in zip(other_labels, other_values):
                        if (label is not None and str(label).strip() != "" and 
                            value is not None and str(value).strip() != "" and 
                            str(value).strip() != "0"):
                            filtered_data.append((label, value))
                    
                    if filtered_data:
                        filtered_labels, filtered_values = zip(*filtered_data)
                    else:
                        filtered_labels, filtered_values = [], []
                    
                    # Create individual bars (not stacked)
                    if filtered_labels and filtered_values:
                        # Use individual colors for each bar
                        bar_colors = []
                        for i in range(len(filtered_labels)):
                            if other_colors and i < len(other_colors):
                                bar_colors.append(other_colors[i])
                            elif colors and i < len(colors):
                                bar_colors.append(colors[i])
                            else:
                                bar_colors.append('#1f77b4')  # Default blue
                        
                        bars = ax2.bar(range(len(filtered_labels)), filtered_values, color=bar_colors, alpha=0.7)
                    else:
                        bars = []
                    
                    expanded_segment = chart_meta.get("expanded_segment", "Other")
                    ax2.set_title(f"Breakdown of '{expanded_segment}'", fontsize=font_size or 12, weight='bold')
                    # Use proper Y-axis label from configuration
                    y_axis_title = chart_meta.get("y_axis_title", "Value")
                    ax2.set_ylabel(y_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set X-axis label
                    x_axis_title = chart_meta.get("x_axis_title", "Categories")
                    ax2.set_xlabel(x_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set x-tick labels for filtered data
                    if filtered_labels:
                        ax2.set_xticks(range(len(filtered_labels)))
                        ax2.set_xticklabels(filtered_labels, rotation=0)
                    # Add data labels with proper formatting
                    value_format = chart_meta.get("value_format", ".2f")
                    data_label_font_size = chart_meta.get("data_label_font_size", 10)
                    data_label_color = chart_meta.get("data_label_color", "#000000")
                    for bar, v in zip(bars, filtered_values):
                        if value_format == "value":
                            formatted_value = str(v)
                        elif value_format == ".2f":
                            formatted_value = f"{v:.2f}"
                        elif value_format == ".1f":
                            formatted_value = f"{v:.1f}"
                        elif value_format == ".0f":
                            formatted_value = f"{v:.0f}"
                        else:
                            formatted_value = f"{v:{value_format}}"
                        ax2.text(bar.get_x() + bar.get_width()/2, v, formatted_value, ha='center', va='bottom', fontweight='bold', fontsize=data_label_font_size, color=data_label_color)

                else:
                    # Bar, line, area charts
                    mpl_figsize = figsize if figsize else (10, 6)
                    # current_app.logger.debug(f"Applied Matplotlib figsize: {mpl_figsize}")
                    fig_mpl, ax1 = plt.subplots(figsize=mpl_figsize, dpi=200)
                    ax2 = ax1.twinx()
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)

                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors):
                            color = colors[i]

                        y_vals = series.get("values")
                        value_range = series.get("value_range")
                        if value_range:
                            # Check if value_range is already extracted (list) or still a string
                            if isinstance(value_range, list):
                                y_vals = value_range
                            else:
                                y_vals = extract_values_from_range(value_range)

                        # Generic chart type handling for Matplotlib
                        
                        mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        if mpl_chart_type == "bar":
                            if chart_type == "stacked_column":
                                # For stacked column, use bottom parameter
                                if i == 0:
                                    ax1.bar(x_values, y_vals, label=label, color=color, alpha=0.7)
                                    bottom_vals = y_vals
                                else:
                                    ax1.bar(x_values, y_vals, bottom=bottom_vals, label=label, color=color, alpha=0.7)
                                    bottom_vals = [sum(x) for x in zip(bottom_vals, y_vals)]
                            else:
                                if isinstance(color, list):
                                    for j, val in enumerate(y_vals):
                                        bar_color = color[j % len(color)]
                                        ax1.bar(x_values[j], val, color=bar_color, alpha=0.7, label=label if j == 0 else "")
                                else:
                                    ax1.bar(x_values, y_vals, label=label, color=color, alpha=0.7)
                                    
                        elif mpl_chart_type == "barh":
                            # Horizontal bar chart
                            if isinstance(color, list):
                                for j, val in enumerate(y_vals):
                                    bar_color = color[j % len(color)]
                                    ax1.barh(x_values[j], val, color=bar_color, alpha=0.7, label=label if j == 0 else "")
                            else:
                                ax1.barh(x_values, y_vals, label=label, color=color, alpha=0.7)
                                
                        elif mpl_chart_type == "plot":
                            # Line chart
                            marker = 'o' if series_type == "scatter_line" else None
                            if ax2:
                                ax2.plot(x_values, y_vals, label=label, color=color, marker=marker, linewidth=2)
                            else:
                                ax1.plot(x_values, y_vals, label=label, color=color, marker=marker, linewidth=2)
                            
                        elif mpl_chart_type == "scatter":
                            # Scatter plot and Bubble chart
                            if series_type == "bubble":
                                # Enhanced bubble chart with better styling
                                # Creating bubble chart for series: {label}
                                
                                # Get sizes from the series data structure
                                sizes = series.get("size", [20] * len(y_vals))
                                # Bubble chart data processed
                                
                                # Ensure all arrays have the same length
                                min_length = min(len(x_values), len(y_vals), len(sizes))
                                if min_length < len(x_values) or min_length < len(y_vals) or min_length < len(sizes):
                                    current_app.logger.warning(f"⚠️ Array length mismatch! Truncating to {min_length}")
                                    x_values = x_values[:min_length]
                                    y_vals = y_vals[:min_length]
                                    sizes = sizes[:min_length]
                                    if isinstance(color, list):
                                        color = color[:min_length]
                                
                                # Scale sizes for better visual impact (bubble charts need much larger sizes)
                                scaled_sizes = [s * 20 for s in sizes]  # Much larger scale for better visibility
                                # Scaled sizes calculated
                                
                                # Get colors - support both single color and color list
                                bubble_colors = color
                                if isinstance(color, list):
                                    # Use provided colors
                                    bubble_colors = color
                                elif color:
                                    # Single color - create gradient based on size
                                    import matplotlib.cm as cm
                                    import numpy as np
                                    size_array = np.array(sizes)
                                    normalized_sizes = (size_array - size_array.min()) / (size_array.max() - size_array.min() + 1e-8)
                                    cmap = cm.viridis if color == 'auto' else cm.get_cmap('viridis')
                                    bubble_colors = [cmap(norm_size) for norm_size in normalized_sizes]
                                else:
                                    # Default color scheme
                                    bubble_colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F']
                                
                                # Create bubble chart with enhanced styling
                                scatter = ax1.scatter(x_values, y_vals, 
                                                     s=scaled_sizes, 
                                                     c=bubble_colors, 
                                                     alpha=0.8,  # Increased opacity
                                                     edgecolors='white',  # White borders
                                                     linewidth=2,  # Border thickness
                                                     label=label,
                                                     zorder=3)  # Higher z-order for better layering
                                
                                # Add subtle shadow effect for depth
                                ax1.scatter(x_values, y_vals, 
                                          s=scaled_sizes, 
                                          c='black', 
                                          alpha=0.1, 
                                          zorder=1)
                                
                                # Add data labels for bubble chart if enabled
                                if show_data_labels:
                                    for j, (x, y, size_val) in enumerate(zip(x_values, y_vals, sizes)):
                                        if j < len(x_values):
                                            # Format size value for label
                                            try:
                                                if value_format == ".1f":
                                                    formatted_size = f"{float(size_val):.1f}"
                                                elif value_format == ".0f":
                                                    formatted_size = f"{float(size_val):.0f}"
                                                else:
                                                    formatted_size = f"{float(size_val):.0f}"
                                            except:
                                                formatted_size = str(size_val)
                                            
                                            # Add size label on bubble with better positioning
                                            label_color = data_label_color or '#000000'
                                            ax1.text(x, y, formatted_size, 
                                                    ha='center', va='center', 
                                                    fontsize=data_label_font_size or 10,
                                                    color=label_color,
                                                    fontweight='bold',
                                                    zorder=5)  # Higher z-order to ensure visibility
                            else:
                                # Enhanced scatter plot with custom styling
                                # Extract marker properties from series
                                marker_size = series.get("marker", {}).get("size", 50)
                                marker_color = series.get("marker", {}).get("color", color)
                                marker_symbol = series.get("marker", {}).get("symbol", "o")
                                marker_opacity = series.get("marker", {}).get("opacity", 0.8)
                                text_labels = series.get("text", [])
                                text_position = series.get("textposition", "top center")
                                
                                # Handle marker size (can be single value or list)
                                if isinstance(marker_size, list):
                                    sizes = marker_size
                                else:
                                    sizes = [marker_size] * len(y_vals)
                                
                                # Scale sizes for better visibility
                                scaled_sizes = [s * 2 for s in sizes]
                                
                                # Create enhanced scatter plot
                                scatter = ax1.scatter(x_values, y_vals, 
                                                     s=scaled_sizes, 
                                                     c=marker_color, 
                                                     alpha=marker_opacity,
                                                     edgecolors='white',
                                                     linewidth=1.5,
                                                     label=label,
                                                     zorder=3)
                                
                                # Add subtle shadow effect for depth
                                ax1.scatter(x_values, y_vals, 
                                          s=scaled_sizes, 
                                          c='black', 
                                          alpha=0.1, 
                                          zorder=1)
                                
                                # Add data labels if text is provided
                                if text_labels and len(text_labels) == len(y_vals):
                                    for j, (x, y, text) in enumerate(zip(x_values, y_vals, text_labels)):
                                        if j < len(text_labels):
                                            # Determine label position based on textposition
                                            if "top" in text_position:
                                                y_offset = 2
                                            elif "bottom" in text_position:
                                                y_offset = -2
                                            else:
                                                y_offset = 0
                                            
                                            if "center" in text_position:
                                                ha = 'center'
                                            elif "left" in text_position:
                                                ha = 'left'
                                            elif "right" in text_position:
                                                ha = 'right'
                                            else:
                                                ha = 'center'
                                            
                                            # Add text label
                                            label_color = data_label_color or '#000000'
                                            ax1.text(x, y + y_offset, str(text), 
                                                    ha=ha, va='bottom' if y_offset > 0 else 'top',
                                                    fontsize=data_label_font_size or 10,
                                                    color=label_color,
                                                    fontweight='bold',
                                                    bbox=dict(boxstyle="round,pad=0.3", 
                                                             facecolor='white', 
                                                             alpha=0.9,
                                                             edgecolor='gray',
                                                             linewidth=0.5))
                                
                        elif mpl_chart_type == "fill_between":
                            # Enhanced Area chart
                            # current_app.logger.debug(f"Processing area chart for series: {label}")
                            # current_app.logger.debug(f"X values: {x_vals}")
                            # current_app.logger.debug(f"Y values: {y_vals}")
                            
                            # Extract area-specific properties from series
                            fill_type = series.get("fill", "tozeroy")
                            line_color = series.get("line", {}).get("color", color)
                            line_width = series.get("line", {}).get("width", 2)
                            marker_symbol = series.get("marker", {}).get("symbol", "o")
                            marker_size = series.get("marker", {}).get("size", 6)
                            marker_color = series.get("marker", {}).get("color", line_color)
                            area_opacity = series.get("opacity", 0.6)
                            text_labels = series.get("text", [])
                            text_position = series.get("textposition", "top center")
                            
                            # Create area fill based on fill type
                            if fill_type == "tozeroy":
                                # Fill from zero to y values
                                ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                            elif fill_type == "tonexty":
                                # Fill to next y values (for stacked areas)
                                if i == 0:
                                    # First series - fill from zero
                                    ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                                    bottom_vals = y_vals
                                else:
                                    # Subsequent series - fill from previous series
                                    ax1.fill_between(x_vals, bottom_vals, [b + y for b, y in zip(bottom_vals, y_vals)], 
                                                   alpha=area_opacity, label=label, color=line_color)
                                    bottom_vals = [b + y for b, y in zip(bottom_vals, y_vals)]
                            else:
                                # Default fill
                                ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                            
                            # Add line on top of area
                            ax1.plot(x_vals, y_vals, color=line_color, linewidth=line_width, zorder=3)
                            
                            # Add markers if specified
                            if marker_symbol != "none":
                                ax1.scatter(x_vals, y_vals, color=marker_color, s=marker_size*20, 
                                          zorder=4, edgecolors='white', linewidth=1)
                            
                            # Add data labels if text is provided
                            if text_labels and len(text_labels) == len(y_vals):
                                for j, (x, y, text) in enumerate(zip(x_vals, y_vals, text_labels)):
                                    if j < len(text_labels):
                                        # Determine label position based on textposition
                                        if "top" in text_position:
                                            y_offset = 5
                                        elif "bottom" in text_position:
                                            y_offset = -5
                                        else:
                                            y_offset = 0
                                        
                                        if "center" in text_position:
                                            ha = 'center'
                                        elif "left" in text_position:
                                            ha = 'left'
                                        elif "right" in text_position:
                                            ha = 'right'
                                        else:
                                            ha = 'center'
                                        
                                        # Add text label
                                        label_color = data_label_color or '#000000'
                                        ax1.text(x, y + y_offset, str(text), 
                                               ha=ha, va='bottom' if y_offset > 0 else 'top',
                                               fontsize=data_label_font_size or 10,
                                               color=label_color,
                                               fontweight='bold',
                                               bbox=dict(boxstyle="round,pad=0.3", 
                                                        facecolor='white', 
                                                        alpha=0.9,
                                                        edgecolor='gray',
                                                        linewidth=0.5),
                                               zorder=5)
                            
                        elif mpl_chart_type == "hist":
                            # Histogram
                            ax1.hist(y_vals, bins=10, label=label, color=color, alpha=0.7)
                            
                        elif mpl_chart_type == "boxplot":
                            # Box plot
                            ax1.boxplot(y_vals, labels=[label], patch_artist=True)
                            if color:
                                ax1.findobj(plt.matplotlib.patches.Patch)[-1].set_facecolor(color)
                                
                        elif mpl_chart_type == "violinplot":
                            # Violin plot
                            ax1.violinplot(y_vals, positions=[i])
                            
                        elif mpl_chart_type == "imshow":
                            # Enhanced heatmap implementation
                            heatmap_data = series.get("z", series.get("values", []))
                            # current_app.logger.debug(f"Heatmap data found: {heatmap_data}")
                            
                            # Ensure we have valid heatmap data
                            if not heatmap_data or len(heatmap_data) == 0:
                                current_app.logger.warning(f"⚠️ No heatmap data found in series: {series}")
                                # Create a default heatmap for testing
                                heatmap_data = [[1, 0, 1, 1], [1, 1, 1, 0], [0, 1, 1, 1]]
                                # current_app.logger.debug(f"Using default heatmap data: {heatmap_data}")
                            
                            # Ensure heatmap_data is a 2D array
                            if isinstance(heatmap_data[0], (int, float)):
                                # Flat list - reshape based on x_axis length
                                cols = len(x_values) if x_values else 4
                                rows = len(heatmap_data) // cols
                                if len(heatmap_data) % cols != 0:
                                    rows += 1
                                # Pad with zeros if needed
                                padded_data = heatmap_data + [0] * (rows * cols - len(heatmap_data))
                                heatmap_data = [padded_data[i:i+cols] for i in range(0, len(padded_data), cols)]
                                # current_app.logger.debug(f"Reshaped heatmap data: {heatmap_data}")
                                # current_app.logger.debug(f"X values: {x_values}")
                                # current_app.logger.debug(f"Cols: {cols}, Rows: {rows}")
                            
                            # Get colorscale from series or use default
                            colorscale = series.get("colorscale", "RdYlGn")
                            # current_app.logger.debug(f"Using colorscale: {colorscale}")
                            
                            # Create heatmap with enhanced styling
                            im = ax1.imshow(heatmap_data, cmap=colorscale, aspect='auto', 
                                           interpolation='nearest', alpha=0.8)
                            
                            # Set axis labels
                            if x_values:
                                ax1.set_xticks(range(len(x_values)))
                                ax1.set_xticklabels(x_values, rotation=0, ha='center')
                            
                            # Create employee labels (Y-axis) - use series labels if available
                            y_labels = series.get("y", [])
                            if not y_labels:
                                num_employees = len(heatmap_data)
                                y_labels = [f"Employee {i+1}" for i in range(num_employees)]
                            
                            ax1.set_yticks(range(len(y_labels)))
                            ax1.set_yticklabels(y_labels)
                            
                            # Add colorbar
                            cbar = plt.colorbar(im, ax=ax1, shrink=0.8)
                            cbar.set_label(series.get("name", "Value"), rotation=270, labelpad=15)
                            
                            # Add text annotations on heatmap cells
                            for i in range(len(heatmap_data)):
                                for j in range(len(heatmap_data[0])):
                                    value = heatmap_data[i][j]
                                    # Determine text color based on background
                                    if colorscale == "RdYlGn":
                                        text_color = 'white' if value > 0.5 else 'black'
                                    else:
                                        text_color = 'white' if value == 1 else 'black'
                                    
                                    # Show actual value or custom text
                                    if "text" in series and i < len(series["text"]) and j < len(series["text"][i]):
                                        cell_text = str(series["text"][i][j])
                                    else:
                                        cell_text = str(value)
                                    
                                    ax1.text(j, i, cell_text, 
                                           ha='center', va='center', color=text_color, 
                                           fontweight='bold', fontsize=10)
                            
                            # Set title
                            ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                            
                            # Remove grid for cleaner look
                            ax1.grid(False)
                            
                            # Set axis labels
                            ax1.set_xlabel(chart_meta.get("x_label", "Week"), fontsize=label_fontsize if 'label_fontsize' in locals() else 12)
                            ax1.set_ylabel(chart_meta.get("primary_y_label", "Employee"), fontsize=label_fontsize if 'label_fontsize' in locals() else 12)
                            
                            # Skip other chart processing for heatmap
                            continue
                                
                        elif mpl_chart_type == "contour":
                            # Contour plot (simplified)
                            if len(y_vals) > 0:
                                # Create a simple 2D array for contour
                                contour_data = [y_vals] if len(y_vals) > 0 else [[0]]
                                ax1.contour(contour_data)
                                
                        else:
                            # Fallback to scatter for unknown types
                            current_app.logger.warning(f"⚠️ Unknown matplotlib chart type '{series_type}', falling back to scatter")
                            ax1.scatter(x_values, y_vals, label=label, color=color, alpha=0.7)

                    # Add data labels to Matplotlib chart if enabled (skip for area charts as they have custom label handling)
                    if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color) and chart_type != "area":
                        # current_app.logger.debug(f"Adding data labels to Matplotlib chart")
                        for i, series in enumerate(series_data):
                            series_type = series.get("type", "bar").lower()
                            y_vals = series.get("values")
                            value_range = series.get("value_range")
                            if value_range:
                                if isinstance(value_range, list):
                                    y_vals = value_range
                                else:
                                    y_vals = extract_values_from_range(value_range)
                            
                            if y_vals:
                                # Determine format to use
                                format_to_use = value_format if value_format else data_label_format
                                if not format_to_use:
                                    format_to_use = ".1f"  # Default format
                                
                                # Add data labels based on chart type
                                if series_type == "bar":
                                    for j, val in enumerate(y_vals):
                                        if j < len(x_values):
                                            # Format the value
                                            try:
                                                if format_to_use == ".1f":
                                                    formatted_val = f"{float(val):.1f}"
                                                elif format_to_use == ".0f":
                                                    formatted_val = f"{float(val):.0f}"
                                                elif format_to_use == ".0%":
                                                    formatted_val = f"{float(val):.0%}"
                                                else:
                                                    formatted_val = f"{float(val):{format_to_use}}"
                                            except:
                                                formatted_val = str(val)
                                            
                                            # Add text label on top of bar (ENHANCED VERSION)
                                            label_color = data_label_color or '#000000'
                                            ax1.text(j, val, formatted_val, 
                                                    ha='center', va='bottom', 
                                                    fontsize=data_label_font_size or int(title_fontsize * 0.9),  # Improved scaling
                                                    color=label_color,
                                                    fontweight='bold',
                                                    bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9))
                                            # current_app.logger.debug(f"Added bar data label with color: {label_color}")
                                
                                elif series_type == "line":
                                    for j, val in enumerate(y_vals):
                                        if j < len(x_values):
                                            # For line charts, use secondary y-axis format if available
                                            line_format = secondary_y_axis_format if secondary_y_axis_format else format_to_use
                                            if not line_format:
                                                line_format = ".1f"  # Default format
                                            
                                            # Format the value
                                            try:
                                                if line_format == ".1f":
                                                    formatted_val = f"{float(val):.1f}"
                                                elif line_format == ".0f":
                                                    formatted_val = f"{float(val):.0f}"
                                                elif line_format == ".0%":
                                                    formatted_val = f"{float(val):.0%}"
                                                elif line_format == ".1%":
                                                    formatted_val = f"{float(val):.1%}"
                                                else:
                                                    formatted_val = f"{float(val):{line_format}}"
                                            except:
                                                formatted_val = str(val)
                                            
                                            # Add text label above line point (ENHANCED VERSION)
                                            label_color = data_label_color or '#000000'
                                            ax2.text(j, val, formatted_val, 
                                                    ha='center', va='bottom', 
                                                    fontsize=data_label_font_size or int(title_fontsize * 0.9),  # Improved scaling
                                                    color=label_color,
                                                    fontweight='bold',
                                                    bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9))
                                            # current_app.logger.debug(f"Added line data label with color: {label_color}")

                                            # Set labels and styling
                        if chart_type != "pie":
                            # Improved font size scaling for Matplotlib
                            title_fontsize = font_size or 52  # Increased default title size
                            label_fontsize = int((font_size or 52) * 0.9)  # Increased relative size for labels
                            
                            # Apply axis label distances using labelpad parameter
                            x_labelpad = x_axis_label_distance / 10.0 if x_axis_label_distance else 5.0  # Convert to points
                            y_labelpad = y_axis_label_distance / 10.0 if y_axis_label_distance else 5.0  # Convert to points
                            
                            ax1.set_xlabel(chart_meta.get("x_label", chart_config.get("x_axis_title", "X")), 
                                         fontsize=label_fontsize, color=font_color, labelpad=x_labelpad)
                            ax1.set_ylabel(chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Primary Y")), 
                                         fontsize=label_fontsize, color=font_color, labelpad=y_labelpad)
                            if "secondary_y_label" in chart_meta or "secondary_y_label" in chart_config:
                                ax2.set_ylabel(chart_meta.get("secondary_y_label", chart_config.get("secondary_y_label", "Secondary Y")), 
                                             fontsize=label_fontsize, color=font_color, labelpad=y_labelpad)

                        # Determine if this is a bubble chart early
                        is_bubble_chart = any(series.get("type", "").lower() == "bubble" for series in series_data)

                        # Set axis tick font size and control tick visibility
                        tick_fontsize = axis_tick_font_size or int(title_fontsize * 0.8)  # Improved tick size calculation
                        if axis_tick_font_size:
                            ax1.tick_params(axis='x', labelsize=axis_tick_font_size, colors=font_color)  # Removed rotation
                            ax1.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                            if ax2 and not is_bubble_chart:
                                ax2.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                        else:
                            ax1.tick_params(axis='x', labelsize=tick_fontsize, colors=font_color)  # Removed rotation
                            ax1.tick_params(axis='y', labelsize=tick_fontsize, colors=font_color)
                            if ax2 and not is_bubble_chart:
                                ax2.tick_params(axis='y', labelsize=tick_fontsize, colors=font_color)
                        
                        # Tick mark control for Matplotlib
                        if show_x_ticks is not None or show_y_ticks is not None:
                            if show_x_ticks is not None:
                                if not show_x_ticks:
                                    ax1.tick_params(axis='x', length=0)  # Hide tick marks
                                    ax1.set_xticklabels([])  # Hide tick labels
                                else:
                                    ax1.tick_params(axis='x', length=5)  # Show tick marks
                            if show_y_ticks is not None:
                                if not show_y_ticks:
                                    ax1.tick_params(axis='y', length=0)  # Hide tick marks
                                    ax1.set_yticklabels([])  # Hide tick labels
                                    if ax2 and not is_bubble_chart:
                                        ax2.tick_params(axis='y', length=0)  # Hide secondary y-axis tick marks
                                        ax2.set_yticklabels([])  # Hide secondary y-axis tick labels
                                else:
                                    ax1.tick_params(axis='y', length=5)  # Show tick marks
                                    if ax2 and not is_bubble_chart:
                                     ax2.tick_params(axis='y', length=5)  # Show secondary y-axis tick marks
                        
                        # Apply X-axis label distance using tick parameters
                        if x_axis_label_distance:
                            # Use tick label padding to control distance
                            ax1.tick_params(axis='x', pad=x_axis_label_distance)
                            # current_app.logger.debug(f"Applied X-axis tick padding: {x_axis_label_distance}")
                        
                        # Apply secondary y-axis formatting for Matplotlib
                        if ax2 and not is_bubble_chart:
                            if secondary_y_axis_format:
                                from matplotlib.ticker import FuncFormatter
                                def percentage_formatter(x, pos):
                                    return f'{x:.0%}'
                                ax2.yaxis.set_major_formatter(FuncFormatter(percentage_formatter))
                        if secondary_y_axis_min_max:
                            # Handle "auto" value for secondary y-axis min/max
                            if secondary_y_axis_min_max == "auto":
                                # Don't set range for auto - let Matplotlib auto-scale
                                pass
                            elif isinstance(secondary_y_axis_min_max, list) and len(secondary_y_axis_min_max) == 2:
                                ax2.set_ylim(secondary_y_axis_min_max)
                            else:
                                current_app.logger.warning(f"Invalid secondary_y_axis_min_max format: {secondary_y_axis_min_max}")
                        elif is_bubble_chart:
                            # current_app.logger.info(f"🎈 Skipping secondary Y-axis formatting for bubble chart")
                            # Explicitly hide secondary Y-axis for bubble charts
                            if ax2:
                                ax2.set_visible(False)
                                # current_app.logger.info(f"🎈 Secondary Y-axis hidden for bubble chart")
                        
                        # Gridlines
                        # current_app.logger.debug(f"Gridlines setting: {show_gridlines}")
                        if show_gridlines:
                            # Map gridline styles to valid Matplotlib linestyles
                            matplotlib_linestyle_map = {
                                "solid": "-",
                                "dashed": "--", 
                                "dash": "--",  # Map 'dash' to '--'
                                "dashdot": "-.",
                                "dotted": ":",
                                "dot": ":",
                                "dotdash": "-."
                            }
                            mapped_linestyle = matplotlib_linestyle_map.get(gridline_style, "--")
                            ax1.grid(True, linestyle=mapped_linestyle, color=gridline_color or '#ccc', alpha=0.6)
                            if ax2 and not is_bubble_chart:
                             ax2.grid(True, linestyle=mapped_linestyle, color=gridline_color or '#ccc', alpha=0.6)
                        else:
                            ax1.grid(False)
                            if ax2:
                             ax2.grid(False)
                        
                        # Apply primary y-axis formatting for Matplotlib
                        if axis_tick_format:
                            from matplotlib.ticker import FuncFormatter
                            if "$" in axis_tick_format:
                                def currency_formatter(x, pos):
                                    return f'${x:,.0f}'
                                ax1.yaxis.set_major_formatter(FuncFormatter(currency_formatter))
                        if y_axis_min_max:
                            # current_app.logger.debug(f"Setting Matplotlib Y-axis range to: {y_axis_min_max}")
                            # Ensure the range is properly applied
                            if isinstance(y_axis_min_max, list) and len(y_axis_min_max) == 2:
                                # Check if the provided range is appropriate for the data
                                all_y_values = []
                                for series in series_data:
                                    y_vals = series.get("values", [])
                                    if y_vals:
                                        all_y_values.extend(y_vals)
                                
                                if all_y_values:
                                    data_min = min(all_y_values)
                                    data_max = max(all_y_values)
                                    data_range = data_max - data_min
                                    
                                    # Check if the provided range is reasonable (within 10x of data range)
                                    provided_range = y_axis_min_max[1] - y_axis_min_max[0]
                                    if provided_range > data_range * 10:
                                        # Auto-calculate appropriate range
                                        padding = data_range * 0.1  # 10% padding
                                        auto_min = max(0, data_min - padding)
                                        auto_max = data_max + padding
                                        ax1.set_ylim(auto_min, auto_max)
                                        # current_app.logger.debug(f"Auto-adjusted Y-axis range from {y_axis_min_max} to {[auto_min, auto_max]} (data range: {data_min}-{data_max})")
                                    else:
                                        # Use provided range
                                        ax1.set_ylim(y_axis_min_max[0], y_axis_min_max[1])
                                        # current_app.logger.debug(f"Applied Y-axis range: {y_axis_min_max[0]} to {y_axis_min_max[1]}")
                                else:
                                    # No data available, use provided range
                                    ax1.set_ylim(y_axis_min_max[0], y_axis_min_max[1])
                                    # current_app.logger.debug(f"Applied Y-axis range: {y_axis_min_max[0]} to {y_axis_min_max[1]}")
                            else:
                                current_app.logger.warning(f"Invalid Y-axis range format: {y_axis_min_max}")
                        
                        # Set title with improved font size
                        ax1.set_title(title, fontsize=title_fontsize, weight='bold', pad=20)
                        
                        # Legend
                        show_legend = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        # current_app.logger.debug(f"Matplotlib legend setting: {show_legend}")
                        
                        # Disable legend for bubble charts to avoid the orange bubble marker
                        if is_bubble_chart:
                            show_legend = False
                            # current_app.logger.info(f"🎈 Disabled legend for bubble chart to avoid extra bubble marker")
                        
                        # Initialize legend_loc outside the if block to fix scope issue
                        legend_loc = 'best'  # default
                        if show_legend:
                            # For dual-axis charts, combine legends from both axes
                            lines1, labels1 = ax1.get_legend_handles_labels()
                            lines2, labels2 = ax2.get_legend_handles_labels()
                            
                            # Map legend position for Matplotlib
                            # current_app.logger.debug(f"Legend position from config: {legend_position}")
                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right"
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                                # current_app.logger.debug(f"Matplotlib legend position mapping: {legend_position} -> {legend_loc}")
                            else:
                                # current_app.logger.debug("No legend position specified, using default 'best'")
                            
                            # Force legend to bottom if specified
                                if legend_position == "bottom":
                                    ax1.legend(lines1 + lines2, labels1 + labels2, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                    # current_app.logger.debug(f"Added legend with {len(lines1 + lines2)} items at forced bottom position")
                                else:
                                    # Skip legend for bubble charts to avoid the orange bubble marker
                                    if is_bubble_chart:
                                        current_app.logger.info(f"🎈 Skipping legend for bubble chart to avoid extra bubble marker")
                                    else:
                                        ax1.legend(lines1 + lines2, labels1 + labels2, loc=legend_loc, fontsize=legend_font_size)
                                        # current_app.logger.debug(f"Added legend with {len(lines1 + lines2)} items at position: {legend_loc}")
                        
                        # Set title with proper font attributes (ENHANCED VERSION)
                        title_fontsize = font_size or 52  # Increased default size
                        if font_color:
                            ax1.set_title(title, fontsize=title_fontsize, weight='bold', color=font_color, 
                                         bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))
                            # current_app.logger.debug(f"Applied title with color: {font_color}, size: {title_fontsize}")
                        else:
                            ax1.set_title(title, fontsize=title_fontsize, weight='bold',
                                         bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))
                        
                        # Apply font color to axis labels (ENHANCED VERSION)
                        x_label_text = chart_meta.get("x_label", "X")
                        primary_y_label_text = chart_meta.get("primary_y_label", "Y")
                        label_fontsize = int(title_fontsize * 0.9)  # Increased relative size
                        
                        # Apply axis label distances using labelpad parameter
                        x_labelpad = x_axis_label_distance / 10.0 if x_axis_label_distance else 5.0  # Convert to points
                        y_labelpad = y_axis_label_distance / 10.0 if y_axis_label_distance else 5.0  # Convert to points
                        
                        if font_color:
                            ax1.set_xlabel(x_label_text, fontsize=label_fontsize, weight='bold', color=font_color,
                                         bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=x_labelpad)
                            ax1.set_ylabel(primary_y_label_text, fontsize=label_fontsize, weight='bold', color=font_color,
                                         bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=y_labelpad)
                            # current_app.logger.debug(f"Applied axis labels with color: {font_color}")
                        else:
                            ax1.set_xlabel(x_label_text, fontsize=label_fontsize, weight='bold',
                                         bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=x_labelpad)
                            ax1.set_ylabel(primary_y_label_text, fontsize=label_fontsize, weight='bold',
                                         bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=y_labelpad)
                        
                        # Apply font size and color to tick labels (ENHANCED VERSION)
                        tick_fontsize = axis_tick_font_size or int(title_fontsize * 0.8)  # Increased relative size
                        if font_color:
                            ax1.tick_params(axis='both', labelsize=tick_fontsize, colors=font_color)
                            # current_app.logger.debug(f"Applied tick labels with color: {font_color}, size: {tick_fontsize}")
                        else:
                            ax1.tick_params(axis='both', labelsize=tick_fontsize)
                        
                        # Tick mark control for Matplotlib (additional section)
                        if show_x_ticks is not None or show_y_ticks is not None:
                            if show_x_ticks is not None:
                                if not show_x_ticks:
                                    ax1.tick_params(axis='x', length=0)  # Hide tick marks
                                    ax1.set_xticklabels([])  # Hide tick labels
                                else:
                                    ax1.tick_params(axis='x', length=5)  # Show tick marks
                            if show_y_ticks is not None:
                                if not show_y_ticks:
                                    ax1.tick_params(axis='y', length=0)  # Hide tick marks
                                    ax1.set_yticklabels([])  # Hide tick labels
                                    if 'ax2' in locals():
                                        ax2.tick_params(axis='y', length=0)  # Hide secondary y-axis tick marks
                                        ax2.set_yticklabels([])  # Hide secondary y-axis tick labels
                                else:
                                    ax1.tick_params(axis='y', length=5)  # Show tick marks
                                    if 'ax2' in locals():
                                        ax2.tick_params(axis='y', length=5)  # Show secondary y-axis tick marks
                        
                        # Also apply font color to secondary y-axis (ENHANCED VERSION)
                        if 'ax2' in locals() and ax2 is not None:
                            # Skip secondary Y-axis for bubble charts
                            is_bubble_chart = any(series.get("type", "").lower() == "bubble" for series in series_data)
                            if not is_bubble_chart:
                                if font_color:
                                    ax2.set_ylabel(chart_meta.get("secondary_y_label", "Secondary Y"), fontsize=label_fontsize, weight='bold', color=font_color,
                                                bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=y_labelpad)
                                    ax2.tick_params(axis='y', labelsize=tick_fontsize, colors=font_color)
                                    # current_app.logger.debug(f"Applied secondary axis with color: {font_color}")
                                else:
                                    ax2.set_ylabel(chart_meta.get("secondary_y_label", "Secondary Y"), fontsize=label_fontsize, weight='bold',
                                                bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.8), labelpad=y_labelpad)
                                ax2.tick_params(axis='y', labelsize=tick_fontsize)
                        else:
                            # current_app.logger.info(f"🎈 Skipping secondary Y-axis for bubble chart")
                
                # Apply axis label distances for Matplotlib
                # current_app.logger.debug(f"X-axis label distance: {x_axis_label_distance}")
                # current_app.logger.debug(f"Y-axis label distance: {y_axis_label_distance}")
                
                            if x_axis_label_distance or y_axis_label_distance:
                                # Get current subplot parameters
                                current_bottom = fig_mpl.subplotpars.bottom
                                current_left = fig_mpl.subplotpars.left
                                
                                if x_axis_label_distance:
                                    # Convert the distance to a fraction of the figure height
                                    # Higher x_axis_label_distance values will push labels further down
                                    adjustment = x_axis_label_distance / 500.0  # Increased conversion factor for more visible effect
                                    fig_mpl.subplots_adjust(bottom=current_bottom - adjustment)
                                    # current_app.logger.debug(f"Applied X-axis adjustment: {adjustment}")
                                
                                if y_axis_label_distance:
                                    # Convert the distance to a fraction of the figure width
                                    adjustment = y_axis_label_distance / 500.0  # Increased conversion factor
                                    fig_mpl.subplots_adjust(left=current_left - adjustment)
                                    # current_app.logger.debug(f"Applied Y-axis adjustment: {adjustment}")
                
                # Apply tight_layout but preserve manual adjustments
                fig_mpl.tight_layout()
                
                # Re-apply manual adjustments after tight_layout with larger effect
                if x_axis_label_distance or y_axis_label_distance:
                    if x_axis_label_distance:
                        adjustment = x_axis_label_distance / 300.0  # Even larger effect after tight_layout
                        fig_mpl.subplots_adjust(bottom=fig_mpl.subplotpars.bottom - adjustment)
                        # current_app.logger.debug(f"Re-applied X-axis adjustment: {adjustment}")
                    
                    if y_axis_label_distance:
                        adjustment = y_axis_label_distance / 300.0  # Even larger effect after tight_layout
                        fig_mpl.subplots_adjust(left=fig_mpl.subplotpars.left - adjustment)
                        # current_app.logger.debug(f"Re-applied Y-axis adjustment: {adjustment}")

                # Add annotations if specified
                if annotations:
                    # current_app.logger.debug(f"Adding {len(annotations)} annotations to chart")
                    
                    # For heatmaps, extract x and y values from series data
                    if chart_type == "heatmap" and series_data:
                        heatmap_x_values = series_data[0].get("x", []) if series_data else []
                        heatmap_y_values = series_data[0].get("y", []) if series_data else []
                    else:
                        heatmap_x_values = []
                        heatmap_y_values = []
                    
                    for annotation in annotations:
                        text = annotation.get("text", "")
                        x_value = annotation.get("x_value", 0)
                        y_value = annotation.get("y_value", 0)
                        
                        # Find x position based on x_value
                        x_pos = 0
                        if chart_type == "heatmap" and heatmap_x_values:
                            # For heatmaps, look in the heatmap x values
                            if x_value in heatmap_x_values:
                                x_pos = heatmap_x_values.index(x_value)
                        elif x_values and x_value in x_values:
                            # For other charts, look in the standard x_values
                            x_pos = x_values.index(x_value)
                        elif isinstance(x_value, str) and x_values:
                            # Try to find string value in x_values
                            try:
                                x_pos = x_values.index(x_value)
                            except ValueError:
                                # If not found, use first position
                                x_pos = 0
                                current_app.logger.warning(f"Annotation x_value '{x_value}' not found in x_values: {x_values}")
                        elif isinstance(x_value, (int, float)) and x_values:
                            # For numeric x_value, use the actual value for bubble charts
                            if chart_type == "bubble":
                                x_pos = x_value  # Use actual x value for bubble charts
                            else:
                                # For other charts, find closest position
                                try:
                                    x_pos = int(x_value)
                                    if x_pos >= len(x_values):
                                        x_pos = len(x_values) - 1
                                except (ValueError, TypeError):
                                    x_pos = 0
                        
                        # Find y position based on y_value
                        y_pos = y_value
                        if chart_type == "heatmap" and heatmap_y_values:
                            # For heatmaps, look in the heatmap y values
                            if y_value in heatmap_y_values:
                                y_pos = heatmap_y_values.index(y_value)
                        
                        # Auto-adjust Y position if it's outside the data range
                        all_y_values = []
                        for series in series_data:
                            y_vals = series.get("values", [])
                            if y_vals:
                                all_y_values.extend(y_vals)
                        
                        if all_y_values:
                            data_min = min(all_y_values)
                            data_max = max(all_y_values)
                            data_range = data_max - data_min
                            
                            # If annotation Y value is way outside data range, adjust it
                            if y_pos > data_max * 2 or y_pos < data_min * 0.5:
                                # Position annotation above the highest data point
                                y_pos = data_max + (data_range * 0.1)  # 10% above max
                                # current_app.logger.debug(f"Auto-adjusted annotation Y position from {y_value} to {y_pos} (data range: {data_min}-{data_max})")
                        
                        # Add annotation text with better positioning for bubble charts
                        if chart_type == "bubble":
                            # For bubble charts, position annotation above the bubble
                            annotation_y_offset = 2000  # Fixed offset for bubble charts
                            ax1.annotate(text, xy=(x_pos, y_pos), xytext=(x_pos, y_pos + annotation_y_offset),
                                       arrowprops=dict(arrowstyle='->', color='red', lw=1.5),
                                       fontsize=12, color='red', weight='bold',
                                       ha='center', va='bottom',
                                       bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9, edgecolor='red'))
                        else:
                            # For other charts, use original positioning
                         ax1.annotate(text, xy=(x_pos, y_pos), xytext=(x_pos, y_pos + (data_range * 0.05) if 'data_range' in locals() else y_pos + 50),
                                   arrowprops=dict(arrowstyle='->', color='red'),
                                   fontsize=12, color='red', weight='bold',
                                   ha='center', va='bottom')
                        # current_app.logger.debug(f"Added annotation: '{text}' at position ({x_pos}, {y_pos})")

                # Apply margin settings to Matplotlib (FIXED VERSION)
                if margin:
                    # current_app.logger.debug(f"Applying margin to Matplotlib: {margin}")
                    # Use figure padding instead of subplot adjustments
                    # Convert margin values to fractions of figure size (more appropriate scaling)
                    # Use a smaller conversion factor to avoid invalid subplot parameters
                    conversion_factor = 1000.0  # Larger denominator for smaller values
                    left_margin_fraction = margin.get("l", 0) / conversion_factor
                    right_margin_fraction = margin.get("r", 0) / conversion_factor
                    top_margin_fraction = margin.get("t", 0) / conversion_factor
                    bottom_margin_fraction = margin.get("b", 0) / conversion_factor
                    
                    # Calculate subplot parameters with validation
                    left_pos = max(0.1, left_margin_fraction)  # Minimum 0.1 to avoid edge
                    right_pos = min(0.9, 1.0 - right_margin_fraction)  # Maximum 0.9 to avoid edge
                    top_pos = min(0.9, 1.0 - top_margin_fraction)  # Maximum 0.9 to avoid edge
                    bottom_pos = max(0.1, bottom_margin_fraction)  # Minimum 0.1 to avoid edge
                    
                    # Ensure left < right and bottom < top
                    if left_pos >= right_pos:
                        left_pos = 0.1
                        right_pos = 0.9
                        current_app.logger.warning(f"Invalid margin: left >= right, using default values")
                    if bottom_pos >= top_pos:
                        bottom_pos = 0.1
                        top_pos = 0.9
                        current_app.logger.warning(f"Invalid margin: bottom >= top, using default values")
                    
                    # Apply margins using figure padding
                    fig_mpl.subplots_adjust(
                        left=left_pos,
                        right=right_pos,
                        top=top_pos,
                        bottom=bottom_pos
                    )
                    # current_app.logger.debug(f"Applied margins (fractions): left={left_pos}, right={right_pos}, top={top_pos}, bottom={bottom_pos}")
                
                # Adjust layout to accommodate legend position
                if show_legend and legend_position == "bottom":
                    # Add extra space at bottom for legend
                    fig_mpl.subplots_adjust(bottom=fig_mpl.subplotpars.bottom + 0.15)
                    # current_app.logger.debug("Added extra bottom space for legend")

                tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                # Use different bbox_inches parameter based on legend position
                if show_legend and legend_position == "bottom":
                    # For bottom legend, use 'tight' but with extra padding
                    plt.savefig(tmpfile.name, bbox_inches='tight', pad_inches=0.3, dpi=200)
                    # current_app.logger.debug(f"Saved Matplotlib chart with bottom legend using extra padding")
                else:
                    # For other positions, use standard tight layout
                    plt.savefig(tmpfile.name, bbox_inches='tight', dpi=200)
                    # Only log legend position if legend_loc is defined (for charts that have legends)
                    if 'legend_loc' in locals():
                         current_app.logger.debug(f"Saved Matplotlib chart with legend at position: {legend_loc}")
                    else:
                        # current_app.logger.debug(f"Saved Matplotlib chart without legend")
                        plt.close(fig_mpl)

                return tmpfile.name

            except Exception as e:
                import traceback
                
                # Create user-friendly error message
                error_type = type(e).__name__
                error_msg = str(e)
                
                # Simplify common error messages
                if "JSONDecodeError" in error_type:
                    user_message = "Invalid JSON format in chart configuration"
                elif "KeyError" in error_type:
                    user_message = f"Missing required field: {error_msg}"
                elif "ValueError" in error_type:
                    user_message = f"Invalid value: {error_msg}"
                elif "IndexError" in error_type:
                    user_message = "Data array is empty or has wrong dimensions"
                elif "TypeError" in error_type:
                    user_message = f"Wrong data type: {error_msg}"
                elif "FileNotFoundError" in error_type:
                    user_message = "Excel file or sheet not found"
                elif "openpyxl" in error_msg.lower():
                    user_message = "Excel file format error - check if file is corrupted"
                elif "pandas" in error_msg.lower():
                    user_message = "Data reading error - check Excel file format"
                else:
                    user_message = error_msg
                
                error_details = {
                    "chart_tag": chart_tag,
                    "error_type": error_type,
                    "user_message": user_message,
                    "technical_message": error_msg,
                    "chart_type": chart_type if 'chart_type' in locals() else "unknown",
                    "data_points": len(series_data) if 'series_data' in locals() else 0,
                    "timestamp": datetime.utcnow().isoformat()
                }
                
                # Simple console logging
                current_app.logger.error(f"❌ Chart '{chart_tag}' failed: {user_message}")
                
                # Store error details for frontend (project-specific)
                if not hasattr(current_app, 'chart_errors'):
                    current_app.chart_errors = {}
                if project_id not in current_app.chart_errors:
                    current_app.chart_errors[project_id] = {}
                current_app.chart_errors[project_id][chart_tag] = error_details
                
                # Also store a simplified version for report generation errors
                if not hasattr(current_app, 'report_generation_errors'):
                    current_app.report_generation_errors = {}
                if project_id not in current_app.report_generation_errors:
                    current_app.report_generation_errors[project_id] = {}
                current_app.report_generation_errors[project_id][chart_tag] = {
                    "error": user_message,
                    "chart_type": chart_type if 'chart_type' in locals() else "unknown",
                    "timestamp": datetime.utcnow().isoformat()
                }
                
                return None

        # Insert charts into paragraphs
        chart_errors = []
        
        # COMPREHENSIVE TEXT REPLACEMENT - Process ALL document elements
        process_entire_document()
        
        # Process charts in paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            full_text = ''.join(run.text for run in para.runs)
            chart_placeholders = re.findall(r"\$\{(section\d+_chart)\}", full_text, flags=re.IGNORECASE)
            for tag in chart_placeholders:
                if tag.lower() in chart_attr_map:
                    try:
                        chart_img = generate_chart({}, tag)
                        if chart_img:
                            para.text = re.sub(rf"\$\{{{tag}\}}", "", para.text, flags=re.IGNORECASE)
                            para.add_run().add_picture(chart_img, width=Inches(5.5))
                        else:
                            # Chart generation failed, add error placeholder
                            error_msg = f"[Chart failed: {tag}]"
                            para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                            
                            # Get the specific error from chart_errors if available
                            specific_error = "Chart could not be generated"
                            if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
                                if tag in current_app.chart_errors[project_id]:
                                    specific_error = current_app.chart_errors[project_id][tag].get('user_message', specific_error)
                            elif hasattr(current_app, 'report_generation_errors') and project_id in current_app.report_generation_errors:
                                if tag in current_app.report_generation_errors[project_id]:
                                    specific_error = current_app.report_generation_errors[project_id][tag].get('error', specific_error)
                            
                            chart_errors.append({
                                "tag": tag,
                                "error": specific_error
                            })
                    except Exception as e:
                        current_app.logger.error(f"⚠️ Failed to insert chart for tag {tag}: {e}")
                        error_msg = f"[Chart failed: {tag}]"
                        para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                        chart_errors.append({
                            "tag": tag,
                            "error": f"Chart insertion failed: {str(e)}"
                        })

        replace_text_in_tables()

        # Insert charts into tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = ''.join(run.text for run in para.runs)
                        chart_placeholders = re.findall(r"\$\{(section\d+_chart)\}", full_text, flags=re.IGNORECASE)
                        for tag in chart_placeholders:
                            if tag.lower() in chart_attr_map:
                                try:
                                    chart_img = generate_chart({}, tag)
                                    if chart_img:
                                        para.text = re.sub(rf"\$\{{{tag}\}}", "", para.text, flags=re.IGNORECASE)
                                        para.add_run().add_picture(chart_img, width=Inches(5.5))
                                    else:
                                        # Chart generation failed, add error placeholder
                                        error_msg = f"[Chart failed: {tag}]"
                                        para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                                        
                                        # Get the specific error from chart_errors if available
                                        specific_error = "Chart could not be generated"
                                        if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
                                            if tag in current_app.chart_errors[project_id]:
                                                specific_error = current_app.chart_errors[project_id][tag].get('user_message', specific_error)
                                        elif hasattr(current_app, 'report_generation_errors') and project_id in current_app.report_generation_errors:
                                            if tag in current_app.report_generation_errors[project_id]:
                                                specific_error = current_app.report_generation_errors[project_id][tag].get('error', specific_error)
                                        
                                        chart_errors.append({
                                            "tag": tag,
                                            "error": specific_error
                                        })
                                except Exception as e:
                                    current_app.logger.error(f"⚠️ Failed to insert chart in table for tag {tag}: {e}")
                                    error_msg = f"[Chart failed: {tag}]"
                                    para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                                    chart_errors.append({
                                        "tag": tag,
                                        "error": f"Chart insertion failed: {str(e)}"
                                    })

        # Save report to temporary location
        import tempfile
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, f'output_report_{project_id}.docx')
        doc.save(output_path)
        current_app.logger.info(f"✅ Report generated successfully")
        
        # Store chart errors for this report generation
        if not hasattr(current_app, 'report_errors'):
            current_app.report_errors = {}
        current_app.report_errors[project_id] = {
            "chart_errors": chart_errors,
            "generated_at": datetime.utcnow().isoformat()
        }
        
        # Clear old chart errors for this project when new report is generated
        if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
            current_app.chart_errors[project_id] = {}
        
        return output_path

    except Exception as e:
        current_app.logger.error(f"❌ Failed to generate report: {e}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return None

# Helper to get absolute path from DB (relative) path

def get_abs_path_from_db_path(db_path):
    # If already absolute, return as is (for backward compatibility)
    if os.path.isabs(db_path):
        return db_path
    # Use the directory containing this file as the root
    return os.path.join(os.path.abspath(os.path.dirname(__file__)), db_path)

@projects_bp.route('/api/projects', methods=['GET'])
@login_required
def get_projects():
    # Access MongoDB via current_app.mongo.db
    projects = list(current_app.mongo.db.projects.find({'user_id': current_user.get_id()}))
    for project in projects:
        project['id'] = str(project['_id'])
        del project['_id'] 
    return jsonify({'projects': projects})

@projects_bp.route('/api/projects', methods=['POST'])
@login_required
def create_project():
    name = request.form.get('name')
    description = request.form.get('description')
    file = request.files.get('file') 

    if not name or not description:
        return jsonify({'error': 'Missing required fields (name or description)'}), 400

    file_path = None
    if file:
        if not allowed_file(file.filename): 
            return jsonify({'error': 'File type not allowed'}), 400
        os.makedirs(UPLOAD_FOLDER, exist_ok=True) 
        filename = secure_filename(file.filename)
        # Store relative path in DB
        file_path = os.path.join('uploads', filename)
        abs_file_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), file_path)
        file.save(abs_file_path)

    project = {
        'name': name,
        'description': description,
        'user_id': current_user.get_id(),
        'file_path': file_path,
        'created_at': datetime.utcnow().isoformat() 
    }
    # Access MongoDB via current_app.mongo.db
    project_id = current_app.mongo.db.projects.insert_one(project).inserted_id
    project['id'] = str(project_id)
    del project['_id']

    return jsonify({'message': 'Project created successfully', 'project': project}), 201

@projects_bp.route('/api/projects/<project_id>/upload_report', methods=['POST'])
@login_required
def upload_report(project_id):
    current_app.logger.info(f"📤 Upload request received for project: {project_id}")
    
    if 'report_file' not in request.files:
        current_app.logger.error(f"❌ No report_file in request.files: {list(request.files.keys())}")
        return jsonify({'error': 'No report file provided'}), 400

    report_file = request.files['report_file']
    current_app.logger.info(f"📁 File received: {report_file.filename}")

    if report_file.filename == '':
        current_app.logger.error(f"❌ Empty filename")
        return jsonify({'error': 'No selected report file'}), 400

    if not allowed_report_file(report_file.filename):
        current_app.logger.error(f"❌ File type not allowed: {report_file.filename}")
        return jsonify({'error': 'Report file type not allowed. Only .xlsx or .csv are accepted.'}), 400

    try:
        project_id_obj = ObjectId(project_id)
        current_app.logger.info(f"✅ Valid project ID: {project_id}")
    except Exception as e:
        current_app.logger.error(f"❌ Invalid project ID: {project_id}, error: {e}")
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        current_app.logger.error(f"❌ Project not found or unauthorized: {project_id}")
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    current_app.logger.info(f"✅ Project found: {project.get('name', 'Unknown')}")

    template_file_path = project.get('file_path')
    abs_template_file_path = get_abs_path_from_db_path(template_file_path)
    current_app.logger.info(f"📄 Template path: {template_file_path}")
    current_app.logger.info(f"📄 Absolute template path: {abs_template_file_path}")
    
    if not template_file_path or not os.path.exists(abs_template_file_path):
        current_app.logger.error(f"❌ Template file not found: {abs_template_file_path}")
        return jsonify({'error': 'Word template file not found for this project. Please upload it during project creation.'}), 400
    
    # Save the uploaded report data file temporarily
    report_data_filename = secure_filename(report_file.filename)
    temp_dir = tempfile.mkdtemp()
    temp_report_data_path = os.path.join(temp_dir, report_data_filename)
    report_file.save(temp_report_data_path)

    # Clear any existing errors for this project before starting new generation
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}

    # Generate the report
    current_app.logger.info(f"🔄 Starting report generation...")
    generated_report_path = _generate_report(project_id, abs_template_file_path, temp_report_data_path)
    
    # Clean up the temporary uploaded report data file and directory
    import shutil
    shutil.rmtree(temp_dir)
    current_app.logger.info(f"🧹 Temporary files cleaned up")

    if generated_report_path:
        current_app.logger.info(f"✅ Report generated successfully: {generated_report_path}")
        # Update project with generated report path
        current_app.mongo.db.projects.update_one(
            {'_id': project_id_obj},
            {'$set': {'generated_report_path': generated_report_path, 'report_generated_at': datetime.utcnow().isoformat()}}
        )
        return jsonify({'message': 'Report generated successfully', 'report_path': generated_report_path}), 200
    else:
        current_app.logger.error(f"❌ Report generation failed")
        return jsonify({'error': 'Failed to generate report'}), 500

@projects_bp.route('/api/reports/<project_id>/download', methods=['GET'])
@login_required
def download_report(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    generated_report_path = project.get('generated_report_path')
    if not generated_report_path or not os.path.exists(generated_report_path):
        return jsonify({'error': 'Generated report not found for this project'}), 404

    # Send file and clean up after download
    try:
        response = send_file(generated_report_path, as_attachment=True)
        
        # Clean up the temporary file after sending
        def cleanup_after_response(response):
            try:
                if os.path.exists(generated_report_path):
                    os.remove(generated_report_path)
                    # Also remove the temp directory if it's empty
                    temp_dir = os.path.dirname(generated_report_path)
                    if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                        os.rmdir(temp_dir)
            except Exception as e:
                current_app.logger.warning(f"⚠️ Failed to cleanup temporary report file: {e}")
            return response
        
        response.call_on_close(lambda: cleanup_after_response(response))
        return response
        
    except Exception as e:
        # Clean up on error too
        try:
            if os.path.exists(generated_report_path):
                os.remove(generated_report_path)
                temp_dir = os.path.dirname(generated_report_path)
                if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                    os.rmdir(temp_dir)
        except:
            pass
        raise e

@projects_bp.route('/api/reports/<chart_filename>/download_html', methods=['GET'])
@login_required
def download_chart_html(chart_filename):
    # Charts are embedded in Word documents, no separate HTML files needed
    return jsonify({'error': 'Chart HTML files are not available - charts are embedded in Word documents'}), 404

@projects_bp.route('/api/reports/batch_reports_<project_id>.zip', methods=['GET'])
@login_required
def download_batch_reports(project_id):
    """Download batch reports ZIP file"""
    try:
        # Validate project ID
        project_id_obj = ObjectId(project_id)
        
        # Check if project exists and belongs to user
        project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
        if not project:
            return jsonify({'error': 'Project not found or unauthorized'}), 404
        
        # Construct the ZIP file path
        zip_filename = f'batch_reports_{project_id}.zip'
        zip_path = os.path.join(tempfile.gettempdir(), zip_filename)
        
        # Check if file exists
        if not os.path.exists(zip_path):
            return jsonify({'error': 'Batch reports file not found. Please regenerate the reports.'}), 404
        
        # Send file
        def cleanup_after_response(response):
            try:
                # Clean up the ZIP file after sending
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                    current_app.logger.info(f"Cleaned up batch reports ZIP: {zip_path}")
            except Exception as e:
                current_app.logger.warning(f"Failed to clean up batch reports ZIP {zip_path}: {e}")
            return response
        
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
        # Add cleanup callback
        response.call_on_close(lambda: cleanup_after_response(response))
        
        current_app.logger.info(f"Downloading batch reports ZIP: {zip_filename}")
        return response
        
    except Exception as e:
        current_app.logger.error(f"Error downloading batch reports: {e}")
        return jsonify({'error': 'Failed to download batch reports'}), 500

@projects_bp.route('/api/projects/<project_id>/chart_errors', methods=['GET'])
@login_required
def get_chart_errors(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Get chart errors for this project
    chart_errors = getattr(current_app, 'chart_errors', {}).get(project_id, {})
    report_errors = getattr(current_app, 'report_errors', {}).get(project_id, {})
    report_generation_errors = getattr(current_app, 'report_generation_errors', {}).get(project_id, {})
    
    # Combine both types of errors
    all_errors = {
        "chart_generation_errors": chart_errors,
        "report_generation_errors": report_errors.get("chart_errors", []),
        "report_generation_errors_detailed": report_generation_errors,
        "report_generated_at": report_errors.get("generated_at")
    }
    
    return jsonify(all_errors)

@projects_bp.route('/api/projects/<project_id>/clear_errors', methods=['POST'])
@login_required
def clear_project_errors(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Clear chart errors for this project
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}
    
    # Clear report errors for this project
    if hasattr(current_app, 'report_errors') and project_id in current_app.report_errors:
        current_app.report_errors[project_id] = {}
    
    return jsonify({'message': 'Project errors cleared successfully'})

@projects_bp.route('/api/projects/<project_id>/upload_zip', methods=['POST'])
@login_required
def upload_zip_and_generate_reports(project_id):
    if 'zip_file' not in request.files:
        return jsonify({'error': 'No zip file provided'}), 400

    zip_file = request.files['zip_file']
    if not zip_file.filename.endswith('.zip'):
        return jsonify({'error': 'Only .zip files are allowed'}), 400

    # Clear any existing errors for this project before starting new generation
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}

    # Prepare temp directories
    temp_dir = tempfile.mkdtemp()
    extracted_dir = os.path.join(temp_dir, 'extracted')
    os.makedirs(extracted_dir, exist_ok=True)

    # Save and extract ZIP
    zip_path = os.path.join(temp_dir, secure_filename(zip_file.filename))
    zip_file.save(zip_path)
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extracted_dir)

    # Find all Excel files
    excel_files = [os.path.join(extracted_dir, f) for f in os.listdir(extracted_dir) if f.endswith('.xlsx') or f.endswith('.xls')]

    # Prepare temporary output folders
    output_folder_name = os.path.join(temp_dir, 'reports_by_name')
    output_folder_code = os.path.join(temp_dir, 'reports_by_code')
    os.makedirs(output_folder_name, exist_ok=True)
    os.makedirs(output_folder_code, exist_ok=True)

    generated_files = []
    total_files = len(excel_files)
    current_app.logger.info(f"Starting batch processing of {total_files} Excel files")
    
    for idx, excel_path in enumerate(excel_files, 1):
        # Extract report name and code from Excel file
        report_name, report_code = extract_report_info_from_excel(excel_path)
        
        current_app.logger.info(f"Processing file {idx}/{total_files}: {report_name} (Code: {report_code})")

        # Generate report
        template_file_path = current_app.mongo.db.projects.find_one({'_id': ObjectId(project_id)})['file_path']
        abs_template_file_path = get_abs_path_from_db_path(template_file_path)
        output_path = _generate_report(f"{project_id}_{idx}", abs_template_file_path, excel_path)
        if output_path:
            # Save in both folders with proper naming
            name_file_path = os.path.join(output_folder_name, f"{report_name}.docx")
            code_file_path = os.path.join(output_folder_code, f"{report_code}.docx")
            
            shutil.copy(output_path, name_file_path)
            shutil.copy(output_path, code_file_path)
            
            generated_files.append({'name': report_name, 'code': report_code})
            current_app.logger.info(f"✅ Successfully generated report {idx}/{total_files}: {report_name} -> {report_code}")
        else:
            current_app.logger.error(f"❌ Failed to generate report {idx}/{total_files}: {report_name}")
        
        # Log progress
        current_app.logger.info(f"Progress: {idx}/{total_files} reports processed")

    # Create zip file in temporary location with both folder structures
    zip_output_path = os.path.join(temp_dir, f'batch_reports_{project_id}.zip')
    with zipfile.ZipFile(zip_output_path, 'w') as zipf:
        # Add files from both folders to maintain the folder structure
        for file_info in generated_files:
            # Add file from reports_by_name folder
            name_file_path = os.path.join(output_folder_name, f"{file_info['name']}.docx")
            if os.path.exists(name_file_path):
                zipf.write(name_file_path, arcname=f"reports_by_name/{file_info['name']}.docx")
            
            # Add file from reports_by_code folder
            code_file_path = os.path.join(output_folder_code, f"{file_info['code']}.docx")
            if os.path.exists(code_file_path):
                zipf.write(code_file_path, arcname=f"reports_by_code/{file_info['code']}.docx")

    # Move zip to a temporary location that will be cleaned up after download
    final_zip_path = os.path.join(tempfile.gettempdir(), f'batch_reports_{project_id}.zip')
    shutil.move(zip_output_path, final_zip_path)
    
    # Clean up temp directory
    shutil.rmtree(temp_dir)

    current_app.logger.info(f"Batch processing complete. Generated {len(generated_files)} out of {total_files} reports")

    return jsonify({
        'message': f'Generated {len(generated_files)} out of {total_files} reports.',
        'download_zip': final_zip_path,
        'reports': generated_files,
        'total_files': total_files,
        'processed_files': len(generated_files),
        'success_rate': f"{len(generated_files)}/{total_files}"
    })

@projects_bp.route('/api/projects/<project_id>', methods=['PUT'])
@login_required
def update_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        return jsonify({'error': 'Invalid project ID'}), 400

    # Check if project exists and belongs to user
    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Get form data
    name = request.form.get('name')
    description = request.form.get('description')
    file = request.files.get('file')

    # Validate required fields
    if not name or not description:
        return jsonify({'error': 'Missing required fields (name or description)'}), 400

    # Prepare update data
    update_data = {
        'name': name,
        'description': description,
        'updated_at': datetime.utcnow().isoformat()
    }

    # Handle file upload if provided
    if file:
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Only .doc or .docx files are accepted.'}), 400
        
        # Delete old file if it exists
        if project.get('file_path'):
            old_file_path = get_abs_path_from_db_path(project['file_path'])
            if os.path.exists(old_file_path):
                try:
                    os.remove(old_file_path)
                except Exception as e:
                    current_app.logger.warning(f"Could not delete old file {old_file_path}: {e}")

        # Save new file
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        filename = secure_filename(file.filename)
        file_path = os.path.join('uploads', filename)
        abs_file_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), file_path)
        file.save(abs_file_path)
        update_data['file_path'] = file_path

    # Update project in database
    result = current_app.mongo.db.projects.update_one(
        {'_id': project_id_obj, 'user_id': current_user.get_id()},
        {'$set': update_data}
    )

    if result.modified_count == 0:
        return jsonify({'error': 'Failed to update project'}), 500

    # Get updated project
    updated_project = current_app.mongo.db.projects.find_one({'_id': project_id_obj})
    updated_project['id'] = str(updated_project['_id'])
    del updated_project['_id']

    return jsonify({'message': 'Project updated successfully', 'project': updated_project})

@projects_bp.route('/api/projects/<project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        return jsonify({'error': 'Invalid project ID'}), 400

    # Check if project exists and belongs to user
    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Delete associated file if it exists
    if project.get('file_path'):
        file_path = get_abs_path_from_db_path(project['file_path'])
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                current_app.logger.warning(f"Could not delete file {file_path}: {e}")

    # Delete project from database
    result = current_app.mongo.db.projects.delete_one({'_id': project_id_obj, 'user_id': current_user.get_id()})

    if result.deleted_count == 0:
        return jsonify({'error': 'Failed to delete project'}), 500

    return jsonify({'message': 'Project deleted successfully'})

@projects_bp.route('/api/projects/<project_id>', methods=['GET'])
@login_required
def get_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        return jsonify({'error': 'Invalid project ID'}), 400

    # Check if project exists and belongs to user
    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    project['id'] = str(project['_id'])
    del project['_id']
    
    return jsonify({'project': project})
